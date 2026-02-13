# -*- coding: utf-8 -*-
import sys, io

if sys.stdout is not None:
    sys.stdout = io.TextIOWrapper(sys.stdout.detach(), encoding='utf-8', errors='replace')
if sys.stderr is not None:
    sys.stderr = io.TextIOWrapper(sys.stderr.detach(), encoding='utf-8', errors='replace')


"""
NSO Excel Analyzer — v1.0
---------------------------------------------
Chức năng chính:
- Nạp nhiều file Excel, tự động ánh xạ Zone_Bx từ DB_VietSub.xlsx
- Lọc dữ liệu theo:
    • TRẠM BIẾN ÁP (so khớp chứa, không phân biệt dấu/hoa–thường, Enter để lọc)
    • U danh định (Uđd)
    • Thời gian (từ ngày – đến ngày) qua checkbox DateEntry
    • Ngưỡng U THẤP (≤ %Uđd) và U CAO (≥ %Uđd), mặc định 95% / 110% (có thể chỉnh)
    • Zone_Bx
- Dashboard phân tích:
    • Thống kê Umin – Utb – Umax (1 chữ số thập phân)
    • Biểu đồ U thực tế (Line / Scatter)
    • 🌡 Heatmap điện áp theo giờ/ngày
    • 📊 Histogram, 📦 Boxplot
- Báo cáo & xuất dữ liệu:
    • Xuất báo cáo phân tích điện áp Zone_Bx (Excel / Word)
    • Xuất danh sách TBA lỗi (chưa map Zone_Bx)
    • Lưu biểu đồ ra PNG
- Quản lý:
    • Ghi nhớ cấu hình & cache tự động
    • Nút Xóa dữ liệu / Xóa toàn bộ cache
- Hỗ trợ:
    • Dashboard hiệu chỉnh TBA lỗi qua webview
    • Nút ❓ Help (hướng dẫn sử dụng & bản quyền)
#=======================ĐÃ SỬA LỖI CHUYỂN ĐỔI EXE CÓ IN RA CONSOLE TIẾNG VIỆT==========================#
Bản quyền phần mềm © 2025 NSO / SuNV
"""

import plotly.graph_objs as go
import plotly.io as pio
import webview  # pip install pywebview
import tempfile
import os

import os, re, sys, json, shutil, subprocess, tempfile, unicodedata
from pathlib import Path
from typing import List, Optional, Dict

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("TkAgg")
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkcalendar import DateEntry

from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage

APP_NAME = "station_gui_ctk_v8_1"
CFG_PATH = os.path.join(Path.home(), f".{APP_NAME}_cfg.json")
CACHE_PATH = os.path.join(Path.home(), f".{APP_NAME}_last.pkl")



def safe_print(*args, **kwargs):
    try:
        text = " ".join(str(a) for a in args)
        sys.stdout.buffer.write((text + "\n").encode("utf-8", "replace"))
        sys.stdout.flush()
    except Exception:
        pass

def get_db_path(filename="DB_VietSub.xlsx"):
    if getattr(sys, 'frozen', False):  # đang chạy từ exe
        base_path = os.path.dirname(sys.executable)
    else:  # đang chạy file .py
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, filename)

# ==================== Helpers ====================
def normalize_cols(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out.columns = [re.sub(r"\s+", " ", str(c)).strip() for c in out.columns]
    drop = []
    for c in out.columns:
        cname = re.sub(r"\s+", "", str(c)).strip().lower()
        if cname in {"stt", "so tt", "sott"}:
            drop.append(c)
    if drop:
        out = out.drop(columns=drop, errors="ignore")
    return out

def _norm_text(s: str) -> str:
    s = str(s).strip().lower()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", s)

def read_excel_all_sheets_xlsx(path: str) -> dict:
    return pd.read_excel(path, sheet_name=None, engine="openpyxl")

def read_excel_all_sheets_xls(path: str) -> dict:
    return pd.read_excel(path, sheet_name=None, engine="xlrd")

def has_soffice() -> bool:
    return shutil.which("soffice") is not None

def ensure_readable_xlsx(path: str, tmp_dir: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".xlsx": return path
    if ext != ".xls": return path
    # đọc trực tiếp .xls; nếu lỗi thì convert
    try:
        _ = read_excel_all_sheets_xls(path)
        return path
    except Exception:
        base = Path(path).stem
        out_xlsx = os.path.join(tmp_dir, base + ".xlsx")
        try:
            if sys.platform.startswith("win"):
                import win32com.client as win32  # pip install pywin32
                excel = win32.gencache.EnsureDispatch("Excel.Application")
                excel.DisplayAlerts = False
                wb = excel.Workbooks.Open(path)
                wb.SaveAs(out_xlsx, FileFormat=51)
                wb.Close(False); excel.Quit()
                if not os.path.exists(out_xlsx):
                    raise RuntimeError("Excel không tạo được file .xlsx")
                return out_xlsx
            else:
                if has_soffice():
                    subprocess.run(["soffice","--headless","--convert-to","xlsx",path,"--outdir",tmp_dir],
                                   check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    out = os.path.join(tmp_dir, Path(path).with_suffix(".xlsx").name)
                    if not os.path.exists(out): raise RuntimeError("LibreOffice không tạo được file .xlsx")
                    return out
                return path
        except Exception as e:
            raise RuntimeError(f"Không thể chuyển .xls: {e}")

def detect_datetime_column(df: pd.DataFrame) -> Optional[str]:
    hints = ["ngay", "thoi gian", "date", "time", "thang", "month", "nam", "year", "ngay gio"]
    for c in df.columns:
        low = str(c).lower()
        if any(h in low for h in hints):
            ser = pd.to_datetime(df[c], errors="coerce", dayfirst=True)
            if ser.notna().sum() > 0: return c
    for c in df.select_dtypes(include=["object"]).columns:
        ser = pd.to_datetime(df[c], errors="coerce", dayfirst=True)
        if ser.notna().sum() > 0: return c
    return None

def pick_voltage_col(df: pd.DataFrame) -> Optional[str]:
    preferred = ["u thực tế", "u thuc te", "utt", "u_tt"]
    hints = ["điện áp","dien ap","voltage","kv","u","ua","ub","uc"]
    for c in df.columns:
        low = str(c).lower()
        if any(p in low for p in preferred) and pd.api.types.is_numeric_dtype(df[c]): return c
    for c in df.columns:
        low = str(c).lower()
        if any(h in low for h in hints) and pd.api.types.is_numeric_dtype(df[c]): return c
    for c in df.columns:
        if pd.api.types.is_numeric_dtype(df[c]) and not str(c).startswith("_"): return c
    return None

def pick_nominal_col(df: pd.DataFrame) -> Optional[str]:
    for c in df.columns:
        low = str(c).lower()
        if ("u danh dinh" in low) or ("u danh định" in low): return c
    for c in df.columns:
        low = str(c).lower()
        if "danh dinh" in low or "danh định" in low: return c
    return None

def detect_station_column(df: pd.DataFrame) -> Optional[str]:
    for c in df.columns:
        if str(c).strip().lower() == "trạm biến áp": return c
    for c in df.columns:
        low = str(c).lower()
        if "tram" in low and ("bien ap" in low or "biến áp" in low): return c
    return df.columns[0] if len(df.columns) else None

def sanitize_sheet_name(name: str, used: set) -> str:
    s = re.sub(r'[\\/*?:\[\]]+', '_', str(name)).strip() or "Sheet"
    s = s[:31]
    base = s; i = 1
    while s in used:
        suf = f"_{i}"
        s = (base[:31-len(suf)] + suf) if len(base)+len(suf) > 31 else base + suf
        i += 1
    used.add(s)
    return s

def combine_from_paths(file_paths: List[str]) -> pd.DataFrame:
    # 1) Khử trùng file_paths (tránh trường hợp dialog trả về trùng, hoặc code gọi lại)
    file_paths = [p for p in file_paths if p]
    file_paths = list(dict.fromkeys(file_paths))  # giữ nguyên thứ tự, bỏ trùng

    all_rows = []
    seen_sig = set()  # chống trùng (file, sheet, signature)

    with tempfile.TemporaryDirectory() as tmpd:
        for f in file_paths:
            readable = ensure_readable_xlsx(f, tmpd)
            book = (
                read_excel_all_sheets_xlsx(readable)
                if readable.lower().endswith(".xlsx")
                else read_excel_all_sheets_xls(readable)
            )

            for sname, df in book.items():
                if df is None or df.shape[0] == 0:
                    continue

                df = normalize_cols(df)

                # 2) Tạo chữ ký nội dung để tránh “cùng 1 sheet bị đọc/append lại”
                #    (nhanh + đủ dùng): (rows, cols, hash header + vài dòng đầu)
                try:
                    head_part = df.head(20).to_csv(index=False)
                except Exception:
                    head_part = str(df.columns.tolist()) + "|" + str(df.shape)

                sig = (os.path.basename(f), str(sname), df.shape[0], df.shape[1], hash(head_part))
                if sig in seen_sig:
                    continue
                seen_sig.add(sig)

                df["_source_file"] = os.path.basename(f)
                df["_sheet"] = sname
                all_rows.append(df)

    if not all_rows:
        return pd.DataFrame()

    combined = pd.concat(all_rows, ignore_index=True, sort=False)

    # 3) drop duplicates theo toàn bộ cột trừ "so tt" (nếu có)
    subset = [c for c in combined.columns if c != "so tt"]
    if subset:
        combined = combined.drop_duplicates(subset=subset, keep="first").reset_index(drop=True)

    combined.insert(0, "so tt", np.arange(1, len(combined) + 1))
    return combined


def detect_compare_column(df: pd.DataFrame) -> Optional[str]:
    """Tìm cột 'SO SÁNH (%)' hoặc tương tự (so sanh, %, etc.)."""
    for c in df.columns:
        low = str(c).lower().strip()
        # nhận diện theo tên
        if ("so sánh" in low) or ("so sanh" in low) or ("%" in low):
            if pd.api.types.is_numeric_dtype(df[c]):
                return c
    return None

# ==================== GUI ====================
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        ctk.set_appearance_mode("system")
        ctk.set_default_color_theme("blue")

        self.title("NSO Voltage Analyzer (v2.4)")
        self.geometry("1320x900")

        self.df: pd.DataFrame = pd.DataFrame()
        self.view_df: pd.DataFrame = pd.DataFrame()
        self.last_dir = os.path.expanduser("~")

        self.voltage_col: Optional[str] = None
        self.compare_col: Optional[str] = None  # cột SO SÁNH (%)
        self.nominal_col: Optional[str] = None
        self.dt_col: Optional[str] = None
        self.zone_filter_str = tk.StringVar(value="-- Tất cả --")

        # --- Filter mode: Auto-apply / Apply ---
        self.auto_apply = ctk.BooleanVar(value=True)

        # --- Zone_Bx multi-select state ---
        self.zone_selected = set()    # set[str] các zone đã chọn
        self.zones_all = []           # list[str] danh sách zone có trong df

        # state vars
        self.chart_mode = tk.StringVar(value="line")
        self.station_text = tk.StringVar()
        self.nominal_val = tk.StringVar()
        self.from_date_str = tk.StringVar()
        self.to_date_str   = tk.StringVar()

        self.use_unom_filter = tk.BooleanVar(value=False)
        self.use_time_filter = tk.BooleanVar(value=False)

        self.use_low_filter  = tk.BooleanVar(value=False)
        self.use_high_filter = tk.BooleanVar(value=False)
        self.low_pct_str  = tk.StringVar(value="95")   # % of Un
        self.high_pct_str = tk.StringVar(value="110")  # % of Un

        # Load cfg
        self.cfg = self._load_cfg()
        if self.cfg.get("geometry"):
            try: self.geometry(self.cfg["geometry"])
            except: pass
        self.last_dir = self.cfg.get("last_dir", self.last_dir)
        self.station_text.set(self.cfg.get("station_text",""))
        self.nominal_val.set(self.cfg.get("nominal_val",""))
        self.from_date_str.set(self.cfg.get("from_date",""))
        self.to_date_str.set(self.cfg.get("to_date",""))
        self.chart_mode.set(self.cfg.get("chart_mode","line"))
        self.use_unom_filter.set(self.cfg.get("use_unom_filter", False))
        self.use_time_filter.set(self.cfg.get("use_time_filter", False))

        self.use_low_filter.set(self.cfg.get("use_low_filter", False))
        self.use_high_filter.set(self.cfg.get("use_high_filter", False))
        self.low_pct_str.set(self.cfg.get("low_pct_str", "95"))
        self.high_pct_str.set(self.cfg.get("high_pct_str", "110"))

        self.voltage_col = self.cfg.get("voltage_col") or None
        self.nominal_col = self.cfg.get("nominal_col") or None

        self._build_gui_modern_card()
        self._try_load_cache()
        self.protocol("WM_DELETE_WINDOW", self._on_close)

    # ---------- Config/cache ----------
    def _load_cfg(self):
        try:
            if os.path.exists(CFG_PATH):
                with open(CFG_PATH,"r",encoding="utf-8") as f:
                    return json.load(f)
        except Exception: pass
        return {}

    def _save_cfg(self):
        from_s = ""
        to_s = ""
        try:
            from_s = self.from_entry.get_date().strftime("%d-%m-%Y")
            to_s = self.to_entry.get_date().strftime("%d-%m-%Y")
        except Exception:
            from_s = self.from_date_str.get()
            to_s = self.to_date_str.get()
        cfg = {
            "geometry": self.winfo_geometry(),
            "last_dir": self.last_dir,
            "voltage_col": self.voltage_col or "",
            "nominal_col": self.nominal_col or "",
            "chart_mode": self.chart_mode.get(),
            "station_text": self.station_text.get(),
            "nominal_val": self.nominal_val.get(),
            "from_date": from_s,
            "to_date": to_s,
            "use_unom_filter": self.use_unom_filter.get(),
            "use_time_filter": self.use_time_filter.get(),
            "use_low_filter": self.use_low_filter.get(),
            "use_high_filter": self.use_high_filter.get(),
            "low_pct_str": self.low_pct_str.get(),
            "high_pct_str": self.high_pct_str.get(),
        }
        try:
            with open(CFG_PATH,"w",encoding="utf-8") as f:
                json.dump(cfg,f,ensure_ascii=False,indent=2)
        except Exception: pass

    def _cache_df(self):
        if not self.df.empty:
            try: self.df.to_pickle(CACHE_PATH)
            except Exception: pass

    def _try_load_cache(self):
        if os.path.exists(CACHE_PATH):
            try:
                df = pd.read_pickle(CACHE_PATH)
                if isinstance(df, pd.DataFrame) and not df.empty:
                    self.df = df.copy()
                    self.view_df = self.df.copy()
                    self._populate_detects()
                    self._refresh_table()
                    self._update_stats_and_chart()
            except Exception as e:
                self._log(f"Không thể nạp cache: {e}")
    def _show_tba_dashboard(self):
        """Alias để gọi dashboard hiệu chỉnh TBA lỗi"""
        return self._show_dashboard_fix_tba_loi()

    # ---------- UI ----------
    def _build_gui_modern_card(self):
        # ---------- Header ----------
        header = ctk.CTkFrame(self, fg_color="#eaf1ff", corner_radius=0, height=58)
        header.pack(side="top", fill="x")
        ctk.CTkLabel(header, text="TOOL ĐIỆN ÁP", font=("Segoe UI", 25, "bold"),
                     fg_color="transparent", text_color="#1a2857").place(x=28, y=12)
        ctk.CTkButton(header, text="❓ Help", width=70, height=36,
                      font=("Segoe UI", 15), command=self._show_help).place(relx=1, x=-18, y=13, anchor="ne")

        # ---------- Main body ----------
        body = ctk.CTkFrame(self, fg_color="#f5f8ff", corner_radius=0)
        body.pack(fill="both", expand=True, padx=0, pady=(0,0))
        body.grid_columnconfigure(0, weight=0, minsize=180)  # Sidebar (button)
        body.grid_columnconfigure(1, weight=1)
        body.grid_rowconfigure(0, weight=0)   # Filter bar
        body.grid_rowconfigure(1, weight=1)   # Main content

        # ---------- Sidebar (CỘT TRÁI) BUTTON ----------
        # ---------- Sidebar (CỘT TRÁI) BUTTON đẹp hơn ----------
        sidebar = ctk.CTkFrame(body, fg_color="#f4f6fb", corner_radius=24)
        sidebar.grid(row=0, column=0, rowspan=2, sticky="nsw", padx=(18,8), pady=(16,14))
        # Optional: Logo nhỏ hoặc tiêu đề phần mềm
        ctk.CTkLabel(sidebar, text="⚡", font=("Segoe UI", 34), fg_color="transparent", text_color="#1976d2").pack(pady=(10, 6))
        ctk.CTkLabel(sidebar, text="Dữ liệu Điện Áp", font=("Segoe UI", 14, "bold"),
                     text_color="#2b3b63", fg_color="transparent").pack(pady=(0, 18))

        # Nút Nạp file
        ctk.CTkButton(
            sidebar, text="  Nạp file", width=160, height=44, corner_radius=18,
            fg_color="#1976d2", hover_color="#1565c0", text_color="#fff",
            font=("Segoe UI", 15, "bold"), image=None,  # hoặc gắn icon PNG ở đây
            anchor="w", command=self._select_and_load
        ).pack(fill="x", padx=18, pady=(2, 13))

        # Nút Xóa
        ctk.CTkButton(
            sidebar, text="  Xóa", width=160, height=44, corner_radius=18,
            fg_color="#ffb4ab", hover_color="#ff6659", text_color="#ad3535",
            font=("Segoe UI", 15, "bold"), anchor="w", command=self._clear_data
        ).pack(fill="x", padx=18, pady=13)

        # Nút Hiệu chỉnh TBA lỗi
        ctk.CTkButton(
            sidebar, text="  Hiệu chỉnh TBA lỗi", width=160, height=44, corner_radius=18,
            fg_color="#ffe59d", hover_color="#ffd54f", text_color="#a16c13",
            font=("Segoe UI", 15, "bold"), anchor="w", command=self._show_dashboard_fix_tba_loi
        ).pack(fill="x", padx=18, pady=13)

        # Nút Dashboard
        ctk.CTkButton(
            sidebar, text="  Dashboard", width=160, height=44, corner_radius=18,
            fg_color="#d0f8ce", hover_color="#a5d6a7", text_color="#257e36",
            font=("Segoe UI", 15, "bold"), anchor="w", command=self._show_dashboard_zone_voltage_report
        ).pack(fill="x", padx=18, pady=13)

        # Nút Xuất TBA lỗi
        ctk.CTkButton(
            sidebar, text="  Xuất TBA lỗi", width=160, height=44, corner_radius=18,
            fg_color="#b3e5fc", hover_color="#81d4fa", text_color="#155f84",
            font=("Segoe UI", 15, "bold"), anchor="w", command=self._export_missing_tba
        ).pack(fill="x", padx=18, pady=(13, 10))

        # ---------- Filter bar (dải ngang trên cùng) ----------
        filter_card = ctk.CTkFrame(body, fg_color="#fff", corner_radius=14)
        filter_card.grid(row=0, column=1, sticky="ew", padx=(2,18), pady=(16,8))
        # Cho dải filter tràn ngang, grid từng cột cho đều hàng
        filter_card.grid_columnconfigure(tuple(range(0, 18)), weight=0)

        ctk.CTkLabel(filter_card, text="BỘ LỌC", font=("Segoe UI", 16, "bold"),
                     text_color="#1a2857").grid(row=0, column=0, sticky="w", padx=14, pady=(10,2), columnspan=12)

        # Tìm trạm
        ctk.CTkLabel(filter_card, text="Tìm trạm:").grid(row=1, column=0, sticky="e", padx=(14,4), pady=8)
        self.entry_search = ctk.CTkEntry(filter_card, width=130, placeholder_text="Tên trạm...", textvariable=self.station_text)
        self.entry_search.grid(row=1, column=1, sticky="w", padx=(0,10), pady=8)
        self.entry_search.bind("<Return>", lambda e: self._apply_filters())

        # Cột vẽ
##        ctk.CTkLabel(filter_card, text="Cột vẽ:").grid(row=1, column=2, sticky="e", padx=(6,4))
##        self.vcol_cmb = ctk.CTkComboBox(filter_card, width=105, values=[""])
##        self.vcol_cmb.grid(row=1, column=3, sticky="w", padx=(0,10))
##        self.vcol_cmb.bind("<<ComboboxSelected>>", lambda e: self._maybe_apply_filters())
        # Không dùng "Cột vẽ" ở bộ lọc nữa -> giữ placeholder để code khác không lỗi
        self.vcol_cmb = ctk.CTkComboBox(filter_card, width=1, values=[""])
        self.vcol_cmb.grid(row=1, column=3)      # đặt đại vào lưới
        self.vcol_cmb.grid_remove()              # ẩn khỏi UI

        # Lọc Udd
        self.use_unom_filter_chk = ctk.CTkCheckBox(filter_card, text="Lọc Udd", variable=self.use_unom_filter,
                                                   command=self._maybe_apply_filters)
        self.use_unom_filter_chk.grid(row=1, column=4, sticky="e", padx=(4,2))
        self.unom_val_cmb = ctk.CTkComboBox(filter_card, width=78, values=[""])
        self.unom_val_cmb.grid(row=1, column=5, sticky="w", padx=(0,10))
        self.unom_val_cmb.bind("<<ComboboxSelected>>", lambda e: self._maybe_apply_filters())

##        # Zone
##        ctk.CTkLabel(filter_card, text="Zone_Bx:").grid(row=1, column=6, sticky="e", padx=(4,4))
##        self.zone_filter_cmb = ctk.CTkComboBox(filter_card, width=112, values=["-- Tất cả --"], variable=self.zone_filter_str)
##        self.zone_filter_cmb.grid(row=1, column=7, sticky="w", padx=(0,10))
##        self.zone_filter_cmb.bind("<<ComboboxSelected>>", lambda e: self._maybe_apply_filters())

        # Zone_Bx (multi-select)
        ctk.CTkLabel(filter_card, text="Zone_Bx:").grid(row=1, column=6, sticky="e", padx=(4,4))

        btn_zone = ctk.CTkButton(
            filter_card, text="Chọn zone…", width=120,
            command=self._open_zone_multiselect
        )
        btn_zone.grid(row=1, column=7, sticky="w", padx=(0,6))

        # badge: hiển thị "Tất cả" / "N zone"
        self.zone_badge_lbl = ctk.CTkLabel(filter_card, text="Tất cả", text_color="#2563eb")
        self.zone_badge_lbl.grid(row=1, column=8, sticky="w", padx=(0,10))

        # Thời gian
        self.use_time_filter_chk = ctk.CTkCheckBox(filter_card, text="Lọc thời gian",
                                                   variable=self.use_time_filter, command=self._maybe_apply_filters)
        self.use_time_filter_chk.grid(row=1, column=9, sticky="e", padx=(6,2))
        from tkcalendar import DateEntry
        self.from_entry = DateEntry(filter_card, width=10, date_pattern="dd-mm-yyyy")
        self.from_entry.grid(row=1, column=10, sticky="w", padx=(0,4))
        self.to_entry = DateEntry(filter_card, width=10, date_pattern="dd-mm-yyyy")
        self.to_entry.grid(row=1, column=11, sticky="w", padx=(0,10))
        try:
            if self.from_date_str.get():
                self.from_entry.set_date(pd.to_datetime(self.from_date_str.get(), dayfirst=True).date())
            if self.to_date_str.get():
                self.to_entry.set_date(pd.to_datetime(self.to_date_str.get(), dayfirst=True).date())
        except Exception:
            pass

        # Ngưỡng thấp/cao
        # Ngưỡng thấp/cao  (SỬA GRID COLUMN ĐỂ KHÔNG ĐÈ NHAU)
        ctk.CTkCheckBox(
            filter_card, text="THẤP (≤ %Udd):", variable=self.use_low_filter,
            command=self._maybe_apply_filters
        ).grid(row=1, column=12, sticky="e", padx=(6,2))

        ctk.CTkEntry(
            filter_card, width=46, textvariable=self.low_pct_str
        ).grid(row=1, column=13, sticky="w", padx=(0,10))

        ctk.CTkCheckBox(
            filter_card, text="U CAO (≥ %Udd):", variable=self.use_high_filter,
            command=self._maybe_apply_filters
        ).grid(row=1, column=14, sticky="e", padx=(6,2))

        ctk.CTkEntry(
            filter_card, width=46, textvariable=self.high_pct_str
        ).grid(row=1, column=15, sticky="w", padx=(0,10))

        # Auto / Apply (tránh lag khi dữ liệu lớn)
        self.auto_apply_chk = ctk.CTkCheckBox(filter_card, text="Auto", variable=self.auto_apply)
        self.auto_apply_chk.grid(row=1, column=16, sticky="e", padx=(6,2))

        ctk.CTkButton(filter_card, text="Apply", width=70, command=self._apply_filters).grid(row=1, column=17, sticky="w", padx=(0,10))

                # ---------- Main content (dữ liệu & phân tích) ----------
        content = ctk.CTkFrame(body, fg_color="#f5f8ff", corner_radius=0)
        content.grid(row=1, column=1, sticky="nsew", padx=(2,18), pady=(0,10))

        # Bảng nhỏ - Dashboard lớn
        content.grid_columnconfigure(0, weight=1, minsize=200)   # cột bảng
        content.grid_columnconfigure(1, weight=5, minsize=980)   # cột dashboard

        content.grid_rowconfigure(0, weight=0)                   # KPI row
        content.grid_rowconfigure(1, weight=1)                   # Main row

        # ===== KPI cards (trên cùng, giống mẫu hiện đại) =====
        kpi_card = ctk.CTkFrame(content, fg_color="#fff", corner_radius=14)
        kpi_card.grid(row=0, column=0, columnspan=2, sticky="ew", padx=(0,0), pady=(8,8))
        self._build_kpi_row(kpi_card)

        # ===== Data table (trái) =====
        data_card = ctk.CTkFrame(content, fg_color="#fff", corner_radius=14)
        data_card.grid(row=1, column=0, sticky="nsew", padx=(0,4), pady=(0,8))

        ctk.CTkLabel(data_card, text="BẢNG DỮ LIỆU", font=("Segoe UI", 16, "bold"),
                     text_color="#1a2857").pack(anchor="w", padx=16, pady=(12, 0))
        self._build_table(data_card)

        # ===== Dashboard (phải) dạng TAB =====
        dash_card = ctk.CTkFrame(content, fg_color="#fff", corner_radius=14)
        dash_card.grid(row=1, column=1, sticky="nsew", padx=(6,0), pady=(0,8))
        ctk.CTkLabel(dash_card, text="DASHBOARD PHÂN TÍCH", font=("Segoe UI", 16, "bold"),
                     text_color="#1a2857").pack(anchor="w", padx=16, pady=(12, 0))

        self._build_dashboard_tabs(dash_card)

        # ---------- Status bar ----------
        self.status_var = tk.StringVar(value="Sẵn sàng.")
        status_bar = ctk.CTkLabel(self, textvariable=self.status_var,
                                  fg_color="#fff", height=30, corner_radius=8,
                                  text_color="#2b2d33", font=("Segoe UI", 13))
        status_bar.pack(side="bottom", fill="x", padx=24, pady=(2, 10))

    def _build_kpi_row(self, parent):
        """KPI row nằm trên cùng (cập nhật theo view_df sau lọc)."""
        self.kpi_vars = {
            "tba": tk.StringVar(value="0"),
            "rows": tk.StringVar(value="0"),
            "umin": tk.StringVar(value="0"),
            "umax": tk.StringVar(value="0"),
            "utb": tk.StringVar(value="0"),
        }

        wrap = ctk.CTkFrame(parent, fg_color="#ffffff", corner_radius=14)
        wrap.pack(fill="x", padx=14, pady=12)
        wrap.grid_columnconfigure((0, 1, 2, 3, 4), weight=1)

        self._kpi_card(wrap, "🧾", "Số dòng",  self.kpi_vars["rows"], "#546e7a", 0)
        self._kpi_card(wrap, "🏭", "Tổng TBA", self.kpi_vars["tba"],  "#1d9bf0", 1)
        self._kpi_card(wrap, "🔻", "Umin",     self.kpi_vars["umin"], "#009688", 2)
        self._kpi_card(wrap, "🔺", "Umax",     self.kpi_vars["umax"], "#1565c0", 3)
        self._kpi_card(wrap, "📊", "Utb",      self.kpi_vars["utb"],  "#7e57c2", 4)

        # update lần đầu
        self._update_kpi_cards()

    def _build_dashboard_tabs(self, parent):
        """Dashboard dạng tab: Tổng quan / Heatmap / Phân phối / Báo cáo."""
        #tabs = ctk.CTkTabview(parent, corner_radius=14)
        tabs = ctk.CTkTabview(parent, corner_radius=14, command=self._on_dashboard_tab_changed)

        tabs.pack(fill="both", expand=True, padx=14, pady=(8, 14))

        tab_overview = tabs.add("Tổng quan")
        tab_heatmap  = tabs.add("Heatmap")
        tab_dist     = tabs.add("Phân phối")
        tab_report   = tabs.add("Báo cáo")

        # ===== TAB: Tổng quan (chart + stats) =====
        top = ctk.CTkFrame(tab_overview, fg_color="transparent")
        top.pack(fill="x", padx=10, pady=(10, 0))

        ctk.CTkLabel(top, text="Chế độ:", font=("Segoe UI", 13)).pack(side="left", padx=(0, 8))
        ctk.CTkRadioButton(top, text="Line", variable=self.chart_mode, value="line",
                           command=self._update_stats_and_chart).pack(side="left", padx=4)
        ctk.CTkRadioButton(top, text="Scatter", variable=self.chart_mode, value="scatter",
                           command=self._update_stats_and_chart).pack(side="left", padx=4)

        ctk.CTkButton(top, text="💾 Lưu PNG", width=110, command=self._export_figure).pack(side="right")

        self.stats_var = tk.StringVar(value="Thống kê: —")
        ctk.CTkLabel(tab_overview, textvariable=self.stats_var, font=("Segoe UI", 12),
                     text_color="#4b5563").pack(anchor="w", padx=12, pady=(6, 0))

        chart_wrap = ctk.CTkFrame(tab_overview, fg_color="transparent")
        chart_wrap.pack(fill="both", expand=True, padx=10, pady=10)

        self.fig, self.ax = plt.subplots(figsize=(6.2, 4), dpi=100)
        self.canvas = FigureCanvasTkAgg(self.fig, master=chart_wrap)
        self.canvas.get_tk_widget().pack(fill="both", expand=True)

        self._draw_chart_empty()
        self._update_stats_and_chart()

        # ===== TAB: Heatmap (VẼ TRỰC TIẾP TRÊN GUI) =====
        ctk.CTkLabel(tab_heatmap, text="Heatmap điện áp theo giờ/ngày",
                     font=("Segoe UI", 14, "bold"), text_color="#1a2857")\
            .pack(anchor="w", padx=12, pady=(14, 6))

        # vùng đặt chart
        self.hm_wrap = ctk.CTkFrame(tab_heatmap, fg_color="transparent")
        self.hm_wrap.pack(fill="both", expand=True, padx=10, pady=10)

        # ===== TAB: Phân phối (VẼ TRỰC TIẾP TRÊN GUI) =====
        ctk.CTkLabel(tab_dist, text="Phân phối U (Histogram + Boxplot)",
                     font=("Segoe UI", 14, "bold"), text_color="#1a2857")\
            .pack(anchor="w", padx=12, pady=(14, 6))

        self.dist_wrap = ctk.CTkFrame(tab_dist, fg_color="transparent")
        self.dist_wrap.pack(fill="both", expand=True, padx=10, pady=10)

        # --- lưu lại tabview + render ngay tab đang chọn ---
        self.dashboard_tabs = tabs
        self.after(80, self._on_dashboard_tab_changed)

        # vẽ tab hiện tại ngay khi mở dashboard
        self.after(50, lambda: self._on_dashboard_tab_changed(self.dashboard_tabs.get()))

        self.report_wrap = ctk.CTkFrame(tab_report, fg_color="transparent")
        self.report_wrap.pack(fill="both", expand=True, padx=10, pady=10)

    def _render_report_zone_charts_on_gui(self):
        """Báo cáo tổng hợp theo Zone_Bx (đẹp/pro):
        - Sort giảm dần
        - Top N
        - Bar ngang dễ đọc
        - Grid nhẹ + annotate số
        - Tiêu đề tổng + mô tả filter
        """
        if getattr(self, "report_wrap", None) is None:
            return

        for w in self.report_wrap.winfo_children():
            w.destroy()

        if self.view_df.empty:
            ctk.CTkLabel(self.report_wrap, text="Chưa có dữ liệu để lập báo cáo.",
                        font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        df = self.view_df.copy()

        if "Zone_Bx" not in df.columns:
            ctk.CTkLabel(self.report_wrap, text="Thiếu cột Zone_Bx để tổng hợp theo vùng.",
                        font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        # tìm cột % so sánh
        pct_col = None
        for c in df.columns:
            if "so sánh" in str(c).lower() and "%" in str(c):
                pct_col = c
                break
        if pct_col is None:
            for c in df.columns:
                if "percent" in str(c).lower() or "pct" in str(c).lower():
                    pct_col = c
                    break
        if pct_col is None:
            ctk.CTkLabel(self.report_wrap, text="Thiếu cột 'SO SÁNH (%)' để xác định vi phạm.",
                        font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        # ngưỡng từ UI
        try:
            low_thr = float(self.low_pct_str.get())
        except Exception:
            low_thr = 95.0
        try:
            high_thr = float(self.high_pct_str.get())
        except Exception:
            high_thr = 110.0

        # chuẩn hóa
        df[pct_col] = pd.to_numeric(df[pct_col], errors="coerce")
        df["Zone_Bx"] = df["Zone_Bx"].astype(str).str.strip()
        df = df.dropna(subset=["Zone_Bx", pct_col])

        if df.empty:
            ctk.CTkLabel(self.report_wrap, text="Không có dữ liệu hợp lệ để tổng hợp.",
                        font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        st_col = detect_station_column(df) or "TRẠM BIẾN ÁP"
        if st_col not in df.columns:
            st_col = None

        low_df = df[df[pct_col] <= low_thr]
        high_df = df[df[pct_col] >= high_thr]

        def _agg(dfx):
            if dfx.empty:
                return pd.Series(dtype="int64"), pd.Series(dtype="int64")
            if st_col:
                n_tba = dfx.groupby("Zone_Bx")[st_col].nunique()
            else:
                n_tba = dfx.groupby("Zone_Bx").size()
            n_times = dfx.groupby("Zone_Bx").size()
            return n_tba, n_times

        low_tba, low_times = _agg(low_df)
        high_tba, high_times = _agg(high_df)

        # --------- UI: thanh điều khiển Top N ----------
        ctrl = ctk.CTkFrame(self.report_wrap, fg_color="transparent")
        ctrl.pack(fill="x", padx=10, pady=(0, 4))

        ctk.CTkLabel(ctrl, text="Top:", font=("Segoe UI", 12)).pack(side="left")
        top_var = ctk.StringVar(value="12")  # mặc định top 12 zone
        top_entry = ctk.CTkEntry(ctrl, width=60, textvariable=top_var)
        top_entry.pack(side="left", padx=(6, 10))

        ctk.CTkLabel(ctrl, text="(zones)", font=("Segoe UI", 12), text_color="#6b7280").pack(side="left")

        # vùng vẽ chart
        chart_holder = ctk.CTkFrame(self.report_wrap, fg_color="transparent")
        chart_holder.pack(fill="both", expand=True, padx=10, pady=(0, 0))

        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

        def _topn(s: pd.Series, n: int):
            if s is None or s.empty:
                return s
            s2 = s.sort_values(ascending=False)
            return s2.head(n)

        def _barh(ax, s: pd.Series, title: str, xlabel: str, color: str):
            if s is None or s.empty:
                ax.text(0.5, 0.5, "Không có dữ liệu", ha="center", va="center")
                ax.set_axis_off()
                return

            y = list(s.index)[::-1]
            x = list(s.values)[::-1]

            ax.barh(y, x, color=color)
            ax.set_title(title, fontsize=12, pad=10)
            #ax.set_xlabel(xlabel)
            ax.set_xlabel(xlabel, labelpad=6)


            # Grid gọn
            ax.grid(axis="x", linestyle="--", alpha=0.25)

            # Chừa biên phải để số không bị cắt
            xmax = max(x) if x else 0
            ax.set_xlim(0, xmax * 1.12 if xmax > 0 else 1)

            # Annotate (đặt sát trong plot, tránh đâm ra ngoài)
            pad = max(0.01 * xmax, 0.2)
            for yi, xi in zip(y, x):
                ax.text(xi + pad, yi, f"{int(xi)}", va="center", fontsize=9, clip_on=False)

            # Tick gọn + chừa khoảng cho nhãn y
            ax.tick_params(axis="y", labelsize=9, pad=6)
            ax.tick_params(axis="x", labelsize=9)

        def _render():
            # clear canvas cũ
            for w in chart_holder.winfo_children():
                w.destroy()

            # đọc top n
            try:
                top_n = int(top_var.get())
                top_n = max(5, min(top_n, 30))
            except Exception:
                top_n = 12

            # chọn top theo “tổng số lần” để phản ánh mức độ nóng (business-friendly)
            # đọc top n
            try:
                top_n = int(top_var.get())
                top_n = max(5, min(top_n, 30))
            except Exception:
                top_n = 12

            # --- 1) LOẠI nan/None/"" khỏi Zone_Bx ngay trong các series ---
            def _clean_zone_index(s: pd.Series) -> pd.Series:
                if s is None or s.empty:
                    return pd.Series(dtype="int64")
                s2 = s.copy()
                s2.index = s2.index.astype(str).str.strip()
                bad = s2.index.str.lower().isin(["nan", "none", ""])
                s2 = s2[~bad]
                return s2

            l_tba0 = _clean_zone_index(low_tba).fillna(0).astype(int) if low_tba is not None else pd.Series(dtype="int64")
            l_tim0 = _clean_zone_index(low_times).fillna(0).astype(int) if low_times is not None else pd.Series(dtype="int64")
            h_tba0 = _clean_zone_index(high_tba).fillna(0).astype(int) if high_tba is not None else pd.Series(dtype="int64")
            h_tim0 = _clean_zone_index(high_times).fillna(0).astype(int) if high_times is not None else pd.Series(dtype="int64")

            # --- 2) CHỌN TOP N theo "điểm nóng" tổng hợp, để đúng Top=5 ---
            # score = low_times + high_times (ưu tiên theo số lần vi phạm)
            score = l_tim0.add(h_tim0, fill_value=0).astype(int)
            score = score[score > 0].sort_values(ascending=False)

            zones = score.head(top_n).index.tolist()

            if not zones:
                ctk.CTkLabel(chart_holder, text="Không có dữ liệu vi phạm theo ngưỡng hiện tại.",
                            font=("Segoe UI", 12), text_color="#6b7280")\
                    .pack(anchor="w", padx=12, pady=12)
                return

            # --- 3) Reindex theo đúng zones, bỏ zone = 0 để biểu đồ gọn ---
            def _re(s0: pd.Series) -> pd.Series:
                if s0 is None or s0.empty:
                    return pd.Series(index=zones, data=[0]*len(zones), dtype="int64")
                s2 = s0.reindex(zones).fillna(0).astype(int)
                return s2

            l_tba = _re(l_tba0)
            l_tim = _re(l_tim0)
            h_tba = _re(h_tba0)
            h_tim = _re(h_tim0)

            # sort hiển thị theo score (để 4 chart cùng thứ tự zone, nhìn "report")
            order = score.reindex(zones).fillna(0).sort_values(ascending=True).index.tolist()  # ascending để barh đẹp (dưới lớn trên nhỏ)
            l_tba = l_tba.reindex(order)
            l_tim = l_tim.reindex(order)
            h_tba = h_tba.reindex(order)
            h_tim = h_tim.reindex(order)



            # figure: 2x2, chiều cao theo số zone (tối thiểu 6.8)
            h = max(6.8, 0.35 * len(zones) + 2.8)
            fig = plt.Figure(figsize=(12.6, h), dpi=100)

            ax1 = fig.add_subplot(221)
            ax2 = fig.add_subplot(222)
            ax3 = fig.add_subplot(223)
            ax4 = fig.add_subplot(224)



            # suptitle: hạ xuống để không đè title subplot
            fig.suptitle(
                "BÁO CÁO TỔNG HỢP VI PHẠM ĐIỆN ÁP THEO ZONE",
                fontsize=15, fontweight="bold", y=0.955
            )

            _barh(ax1, l_tba, f"TBA vi phạm THẤP (≤{low_thr}%)", "Số TBA", "#ef4444")
            _barh(ax2, l_tim, f"Tổng số lần THẤP (≤{low_thr}%)", "Số lần", "#7c3aed")
            _barh(ax3, h_tba, f"TBA vi phạm CAO (≥{high_thr}%)", "Số TBA", "#10b981")
            _barh(ax4, h_tim, f"Tổng số lần CAO (≥{high_thr}%)", "Số lần", "#2563eb")

            # ==== AUTO chừa biên trái vừa đủ: tránh cắt chữ + tránh thừa khoảng trắng ====
            try:
                zones_show = [str(z) for z in order]
                max_len = max((len(z) for z in zones_show), default=8)

                # bắt đầu từ 0.18 (đỡ thừa trắng), tăng nhẹ theo độ dài zone
                left = 0.18 + min(0.12, max(0.0, (max_len - 6) * 0.010))
                left = min(max(left, 0.18), 0.32)  # chặn trong [0.18..0.32]
            except Exception:
                left = 0.22

            # Layout 1 lần duy nhất (BỎ tight_layout để khỏi “đánh nhau”)
            fig.subplots_adjust(
                left=0.15, right=0.9,
                top=0.80, bottom=0.14,      # bottom tăng để không cắt xlabel; top giảm để không đè title
                wspace=0.24, hspace=0.52    # tăng hspace để title + xlabel không đè nhau
            )
            # ---- DỊCH RIÊNG 2 BIỂU ĐỒ BÊN PHẢI SANG PHẢI ----
            for ax in (ax2, ax4):
                pos = ax.get_position()
                ax.set_position([
                    pos.x0 + 0.08,   # 👈 tăng giá trị này nếu muốn sang phải thêm
                    pos.y0,
                    pos.width,
                    pos.height
                ])
            # optional: đồng bộ label (không bắt buộc)
            try:
                fig.align_labels()
            except Exception:
                pass





            canvas = FigureCanvasTkAgg(fig, master=chart_holder)
            canvas.get_tk_widget().pack(fill="both", expand=True)
            canvas.draw()

            self._report_fig = fig
            self._report_canvas = canvas

        # render lần đầu + bind Enter để refresh
        _render()
        top_entry.bind("<Return>", lambda e: _render())



    def _on_dashboard_tab_changed(self, *_):
        """CTkTabview sẽ gọi callback khi đổi tab (thường không truyền tham số).
        Ta tự lấy tên tab hiện tại bằng tabs.get()."""
        try:
            tab_name = self.dashboard_tabs.get()
        except Exception:
            return

        if tab_name == "Heatmap":
            self._render_heatmap_on_gui()
        elif tab_name == "Phân phối":
            self._render_dist_on_gui()
        elif tab_name == "Báo cáo":
            self._render_report_zone_charts_on_gui()


    def _render_heatmap_on_gui(self):
        """Vẽ heatmap trực tiếp vào tab Heatmap."""
        if getattr(self, "hm_wrap", None) is None:
            return

        # clear vùng vẽ cũ
        for w in self.hm_wrap.winfo_children():
            w.destroy()

        if self.view_df.empty:
            ctk.CTkLabel(self.hm_wrap, text="Chưa có dữ liệu để vẽ heatmap.",
                         font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        df = self.view_df.copy()
        dt_col = self.dt_col or detect_datetime_column(df)
        vcol = self.voltage_col

        if not dt_col or dt_col not in df.columns or not vcol or vcol not in df.columns:
            ctk.CTkLabel(self.hm_wrap, text="Thiếu cột thời gian hoặc cột U để vẽ heatmap.",
                         font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        # ghép thêm giờ nếu có cột giờ
        hour_col = None
        for c in df.columns:
            if "giờ" in str(c).lower() or "hour" in str(c).lower():
                hour_col = c
                break

        dt = pd.to_datetime(df[dt_col], errors="coerce", dayfirst=True)
        if hour_col:
            hour_val = pd.to_numeric(df[hour_col], errors="coerce").fillna(0)
            dt = dt + pd.to_timedelta(hour_val, unit="h")

        v = pd.to_numeric(df[vcol], errors="coerce")
        df["__date"] = dt.dt.date
        df["__hour"] = dt.dt.hour
        df["__v"] = v
        tmp = df.dropna(subset=["__date", "__hour", "__v"])

        if tmp.empty:
            ctk.CTkLabel(self.hm_wrap, text="Không có giá trị hợp lệ để vẽ heatmap.",
                         font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        pivot = tmp.pivot_table(index="__hour", columns="__date", values="__v", aggfunc="mean")
        pivot = pivot.reindex(range(24))

        # ---- vẽ matplotlib embed ----
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

        fig = plt.Figure(figsize=(7.0, 4.2), dpi=100)
        ax = fig.add_subplot(111)

        data = pivot.values
        im = ax.imshow(data, aspect="auto", origin="lower")

        ax.set_title("Heatmap U trung bình theo Giờ và Ngày", fontsize=12)
        ax.set_ylabel("Giờ")
        ax.set_xlabel("Ngày")

        # ticks ngày (giảm số tick cho đỡ rối)
        cols = list(pivot.columns)
        if len(cols) > 0:
            step = max(1, len(cols) // 10)
            xticks = list(range(0, len(cols), step))
            ax.set_xticks(xticks)
            ax.set_xticklabels([str(cols[i]) for i in xticks], rotation=45, ha="right", fontsize=9)

        ax.set_yticks(range(0, 24, 2))
        ax.set_yticklabels([str(i) for i in range(0, 24, 2)], fontsize=9)

        fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04, label="U (kV)")
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=self.hm_wrap)
        canvas.get_tk_widget().pack(fill="both", expand=True)
        canvas.draw()

        # giữ tham chiếu tránh bị GC
        self._hm_fig = fig
        self._hm_canvas = canvas

    def _render_dist_on_gui(self):
        """Vẽ histogram + boxplot trực tiếp vào tab Phân phối."""
        if getattr(self, "dist_wrap", None) is None:
            return

        for w in self.dist_wrap.winfo_children():
            w.destroy()

        if self.view_df.empty:
            ctk.CTkLabel(self.dist_wrap, text="Chưa có dữ liệu để vẽ phân phối.",
                         font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        vcol = self.voltage_col
        if not vcol or vcol not in self.view_df.columns:
            ctk.CTkLabel(self.dist_wrap, text="Chưa xác định được cột U để vẽ phân phối.",
                         font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        v = pd.to_numeric(self.view_df[vcol], errors="coerce").dropna()
        if v.empty:
            ctk.CTkLabel(self.dist_wrap, text="Không có giá trị U hợp lệ.",
                         font=("Segoe UI", 12), text_color="#6b7280")\
                .pack(anchor="w", padx=12, pady=12)
            return

        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

        fig = plt.Figure(figsize=(7.0, 4.2), dpi=100)
        ax1 = fig.add_subplot(121)
        ax2 = fig.add_subplot(122)

        ax1.hist(v.values, bins=30)
        ax1.set_title("Histogram U", fontsize=11)
        ax1.set_xlabel("U (kV)")
        ax1.set_ylabel("Số lần")

        ax2.boxplot(v.values, vert=True, showmeans=True)
        ax2.set_title("Boxplot U", fontsize=11)
        ax2.set_ylabel("U (kV)")

        fig.suptitle(f"Phân phối U ({vcol})", fontsize=12)
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=self.dist_wrap)
        canvas.get_tk_widget().pack(fill="both", expand=True)
        canvas.draw()

        self._dist_fig = fig
        self._dist_canvas = canvas


    def _build_table(self, parent):
        import tkinter as tk
        import tkinter.ttk as ttk

        # ==============================
        # OUTER FRAME (CTk - bo góc đẹp)
        # ==============================
        outer = ctk.CTkFrame(parent, corner_radius=12, fg_color="#ffffff")
        outer.pack(fill="both", expand=True, padx=8, pady=(6, 10))

        # ==============================
        # INNER FRAME (tk.Frame - tránh bị che Treeview)
        # ==============================
        wrap = tk.Frame(outer, bg="#ffffff")
        wrap.pack(fill="both", expand=True, padx=8, pady=8)

        wrap.grid_rowconfigure(0, weight=1)
        wrap.grid_columnconfigure(0, weight=1)

        # ==============================
        # STYLE TREEVIEW (PRO)
        # ==============================
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except Exception:
            pass

        style.configure(
            "Pro.Treeview",
            font=("Segoe UI", 10),
            rowheight=24,
            background="#ffffff",
            fieldbackground="#ffffff",
            foreground="#111827",
            borderwidth=0,
            relief="flat",
        )
        style.configure(
            "Pro.Treeview.Heading",
            font=("Segoe UI", 10, "bold"),
            background="#f3f4f6",
            foreground="#111827",
            relief="flat",
            borderwidth=0,
        )
        style.map(
            "Pro.Treeview",
            background=[("selected", "#dbeafe")],
            foreground=[("selected", "#111827")],
        )

        # ==============================
        # TREEVIEW
        # ==============================
        self.table = ttk.Treeview(
            wrap,
            show="headings",
            style="Pro.Treeview",
            selectmode="extended",
        )
        self.table.grid(row=0, column=0, sticky="nsew")

        # ==============================
        # SCROLLBARS
        # ==============================
        vsb = ttk.Scrollbar(wrap, orient="vertical", command=self.table.yview)
        hsb = ttk.Scrollbar(wrap, orient="horizontal", command=self.table.xview)
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")

        self.table.configure(
            yscrollcommand=vsb.set,
            xscrollcommand=hsb.set
        )

        # ==============================
        # ZEBRA ROWS
        # ==============================
        self.table.tag_configure("even", background="#f9fafb")
        self.table.tag_configure("odd", background="#ffffff")

        # ==============================
        # SORT BY CLICK HEADING
        # ==============================
        self._table_sort_state = {}

        def _sort_by(col):
            desc = self._table_sort_state.get(col, False)
            self._table_sort_state[col] = not desc

            data = [(self.table.set(i, col), i) for i in self.table.get_children("")]
            def _key(v):
                try:
                    return float(str(v).replace(",", "").strip())
                except Exception:
                    return str(v).lower()

            data.sort(key=lambda x: _key(x[0]), reverse=desc)

            for idx, (_, iid) in enumerate(data):
                self.table.move(iid, "", idx)

            for idx, iid in enumerate(self.table.get_children("")):
                self.table.item(iid, tags=("even" if idx % 2 == 0 else "odd"))

        def _bind_heading_sort():
            for c in self.table["columns"]:
                self.table.heading(c, command=lambda _c=c: _sort_by(_c))

        self._table_bind_heading_sort = _bind_heading_sort

        # ==============================
        # CTRL + C COPY SELECTED
        # ==============================
        def _copy_selected(event=None):
            sels = self.table.selection()
            if not sels:
                return "break"
            cols = self.table["columns"]
            lines = ["\t".join(cols)]
            for iid in sels:
                vals = [self.table.set(iid, c) for c in cols]
                lines.append("\t".join(vals))
            text = "\n".join(lines)
            self.clipboard_clear()
            self.clipboard_append(text)
            return "break"

        self.table.bind("<Control-c>", _copy_selected)





    # ---------- Actions ----------
    def _log(self, msg: str):
        """Cập nhật trạng thái ngắn gọn ở status bar + in console"""
        try:
            self.status_var.set(msg)
        except Exception:
            pass
        #print(msg)
        safe_print(msg)
        self.update_idletasks()

    def _clear_data(self):
        #"""Xóa toàn bộ dữ liệu hiện tại trong tool"""
        import pandas as pd
        self.df = pd.DataFrame()
        self.view_df = pd.DataFrame()
        self._refresh_table()
        self._update_stats_and_chart()
        self._log("🧹 Đã xóa toàn bộ dữ liệu.")

    def _clear_all(self):
        if not messagebox.askyesno("Xóa dữ liệu", "Bạn có chắc muốn xóa toàn bộ dữ liệu đã nạp và cache?"):
            return
        self.df = pd.DataFrame(); self.view_df = pd.DataFrame()
        try:
            if os.path.exists(CACHE_PATH): os.remove(CACHE_PATH)
        except Exception: pass
        try: self.table.delete(*self.table.get_children())
        except Exception: pass
        self._draw_chart_empty()
        self._log("Đã xóa toàn bộ dữ liệu cũ.")

    def _select_and_load(self):
        initial = self.last_dir if os.path.isdir(self.last_dir) else os.path.expanduser("~")
        paths = filedialog.askopenfilenames(
            title="Chọn (thêm) file Excel",
            initialdir=initial,
            filetypes=[("Excel files", "*.xls *.xlsx"), ("All files", "*.*")]
        )
        if not paths:
            return

        # --- chống trùng đường dẫn ngay từ dialog ---
        paths = [p for p in paths if p]
        paths = list(dict.fromkeys(paths))  # giữ thứ tự, bỏ trùng

        self.last_dir = os.path.dirname(paths[0])

        try:
            # ==========================================================
            # 1) NẠP DỮ LIỆU (CHỈ NẠP 1 LẦN) + CONCAT + DROP DUPLICATES
            # ==========================================================
            new_df = combine_from_paths(list(paths))  # <-- CHỈ GỌI 1 LẦN DUY NHẤT
            if new_df is None or new_df.empty:
                self._log("⚠️ Không có dữ liệu hợp lệ từ các file đã chọn.")
                return

            if self.df is None or self.df.empty:
                combined = new_df.copy()
            else:
                combined = pd.concat([self.df, new_df], ignore_index=True, sort=False)

            # Khử trùng toàn cục theo toàn bộ cột trừ "so tt"
            subset_all = [c for c in combined.columns if c != "so tt"]
            if subset_all:
                combined = combined.drop_duplicates(subset=subset_all, keep="first").reset_index(drop=True)

            # Đánh lại so tt đẹp
            if "so tt" in combined.columns:
                combined = combined.drop(columns=["so tt"], errors="ignore")
            combined.insert(0, "so tt", np.arange(1, len(combined) + 1))

            self.df = combined

            # ==========================================================
            # 2) ÁNH XẠ Zone_Bx (TRIỆT: dọn cột cũ + ép kiểu zone_code)
            # ==========================================================
            try:
                db_path = get_db_path()

                if "TRẠM BIẾN ÁP" not in self.df.columns:
                    self._log("⚠️ Không tìm thấy cột 'TRẠM BIẾN ÁP' để ánh xạ Zone_Bx.")
                elif not os.path.exists(db_path):
                    self._log(f"⚠️ Không tìm thấy file DB_VietSub.xlsx tại: {db_path}")
                else:
                    # ====== MAP Zone_Bx (TRIỆT LỖI zone_code <NA>) ======
                    buses_df = pd.read_excel(db_path, sheet_name="Buses")
                    try:
                        zone_df = pd.read_excel(db_path, sheet_name="Zones")
                    except Exception:
                        zone_df = pd.read_excel(db_path, sheet_name=1)

                    zone_df = zone_df.rename(columns={"zone_name_vi": "Zone_Bx"})

                    # --- helper: dò cột theo danh sách ứng viên ---
                    def _pick_col(df, candidates):
                        cols = {c.lower(): c for c in df.columns}
                        for cand in candidates:
                            if cand in df.columns:
                                return cand
                            if cand.lower() in cols:
                                return cols[cand.lower()]
                        return None

                    # --- helper: ép zone_code an toàn (không rớt NA nếu dữ liệu kiểu "15.0", "15 ") ---
                    def _coerce_zone_code(s):
                        # s: Series
                        x = s.copy()
                        # ưu tiên numeric
                        out = pd.to_numeric(x, errors="coerce")
                        # các giá trị numeric ok
                        ok = out.notna()
                        # phần còn lại: xử lý string "15.0", "15 ", "015"
                        if (~ok).any():
                            t = x[~ok].astype(str).str.strip()
                            t = t.str.replace(".0", "", regex=False)
                            t = t.str.replace(",", ".", regex=False)
                            t2 = pd.to_numeric(t, errors="coerce")
                            out.loc[~ok] = t2
                        return out.astype("Int64")

                    # --- dò đúng tên cột trong DB (tránh DB đặt khác 'zone_code', 'Sym') ---
                    bus_sym_col  = _pick_col(buses_df, ["Sym", "SYM", "sym"])
                    bus_zone_col = _pick_col(buses_df, ["zone_code", "Zone_code", "ZONE_CODE", "zone", "Zone", "ZONE", "zone_id", "Zone_ID", "ZONE_ID"])
                    zone_sym_col  = _pick_col(zone_df, ["Sym", "SYM", "sym"])
                    zone_zone_col = _pick_col(zone_df, ["zone_code", "Zone_code", "ZONE_CODE", "zone", "Zone", "ZONE", "zone_id", "Zone_ID", "ZONE_ID"])

                    if bus_sym_col is None or bus_zone_col is None:
                        self._log(f"⚠️ DB 'Buses' thiếu cột Sym/zone_code (Sym={bus_sym_col}, zone={bus_zone_col}).")
                    else:
                        # chuẩn hóa Sym + zone_code trong buses_df
                        buses_df = buses_df.copy()
                        buses_df[bus_sym_col] = buses_df[bus_sym_col].astype(str).str.strip().str.upper()
                        buses_df[bus_zone_col] = _coerce_zone_code(buses_df[bus_zone_col])

                        if zone_sym_col is None or zone_zone_col is None:
                            self._log(f"⚠️ DB 'Zones' thiếu cột Sym/zone_code (Sym={zone_sym_col}, zone={zone_zone_col}).")
                        else:
                            zone_df = zone_df.copy()
                            zone_df[zone_sym_col] = zone_df[zone_sym_col].astype(str).str.strip().str.upper()
                            zone_df[zone_zone_col] = _coerce_zone_code(zone_df[zone_zone_col])

                            # ===== FIX TRIỆT ĐỂ: zone_code trong Buses là công thức -> pandas đọc ra <NA> =====
                            # Nếu zone_code của Buses bị <NA> hàng loạt (do công thức mất cached result sau khi openpyxl save),
                            # thì suy ra zone_code theo Sym từ sheet Zones (Zones đang là giá trị số ổn định).
                            try:
                                bus_zone_na = buses_df[bus_zone_col].notna().sum()
                                if bus_zone_na == 0 or bus_zone_na < 10:
                                    # map Sym -> zone_code từ Zones
                                    sym2zone = zone_df.set_index(zone_sym_col)[zone_zone_col].to_dict()
                                    buses_df[bus_zone_col] = buses_df[bus_sym_col].map(sym2zone)
                                    buses_df[bus_zone_col] = _coerce_zone_code(buses_df[bus_zone_col])
                                    self._log("ℹ️ zone_code(Buses) là công thức bị mất giá trị -> đã suy ra lại từ sheet Zones.")
                            except Exception as _e:
                                self._log(f"⚠️ Không suy ra được zone_code từ Zones: {_e}")

                            # --- chuẩn hóa key join __jk như code của bạn ---
                            import re, unicodedata

                            def _norm_key(s: str) -> str:
                                s = str(s).strip().lower()
                                s = unicodedata.normalize("NFD", s)
                                s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
                                s = s.replace("đ", "d").replace("Đ", "d")
                                s = re.sub(r"\b\d{2,3}\s*kv\b", " ", s)
                                s = re.sub(r"\b(tba|tram bien ap|nm|tdn|td|xm|nmd|nmdn|nmt|nha may|xi mang|kcn)\b", " ", s)
                                s = re.sub(r"\b\d+[a-z]?\b", " ", s)
                                s = re.sub(r"[,/()\-]", " ", s)
                                s = re.sub(r"\s+", " ", s).strip()
                                return s

                            # Dọn cột cũ để tránh Sym_x/Sym_y / zone_code_x
                            for col in ["__jk", "Sym", "zone_code", "Zone_Bx"]:
                                if col in self.df.columns:
                                    self.df.drop(columns=[col], inplace=True, errors="ignore")

                            buses_df["__jk"] = buses_df["TBA_SCADA"].astype(str).map(_norm_key)
                            self.df["__jk"]  = self.df["TRẠM BIẾN ÁP"].astype(str).map(_norm_key)

                            # --- merge __jk -> Sym, zone_code (đặt tên chuẩn Sym/zone_code) ---
                            bus_map = buses_df[["__jk", bus_sym_col, bus_zone_col]].drop_duplicates(subset=["__jk"]).copy()
                            bus_map = bus_map.rename(columns={bus_sym_col: "Sym", bus_zone_col: "zone_code"})

                            self.df = self.df.merge(bus_map, on="__jk", how="left")

                            # --- merge Sym + zone_code -> Zone_Bx ---
                            zone_map = zone_df[[zone_sym_col, zone_zone_col, "Zone_Bx"]].drop_duplicates(subset=[zone_sym_col, zone_zone_col]).copy()
                            zone_map = zone_map.rename(columns={zone_sym_col: "Sym", zone_zone_col: "zone_code"})

                            self.df["Sym"] = self.df["Sym"].astype(str).str.strip().str.upper()
                            self.df["zone_code"] = _coerce_zone_code(self.df["zone_code"])

                            self.df = self.df.merge(zone_map, on=["Sym", "zone_code"], how="left")

                            # dọn cột tạm
                            self.df.drop(columns=["__jk"], inplace=True, errors="ignore")

                        # Báo cáo gọn
                        if "Zone_Bx" in self.df.columns:
                            missing_rows = self.df[self.df["Zone_Bx"].isna()]
                            if not missing_rows.empty:
                                num_missing = int(missing_rows["TRẠM BIẾN ÁP"].nunique())
                                sample = ", ".join(sorted(missing_rows["TRẠM BIẾN ÁP"].dropna().astype(str).unique()[:5]))
                                self._log(
                                    f"⚠️ Còn {num_missing} trạm chưa ánh xạ Zone_Bx (vd: {sample}…). "
                                    f"Dùng nút '📤 Xuất TBA lỗi' để xuất danh sách chi tiết."
                                )
                            else:
                                self._log("[ok] Đã ánh xạ thành công tất cả TBA sang Zone_Bx.")
            except Exception as e:
                self._log(f"⚠️ Lỗi khi gắn Zone_Bx: {e}")

            # ==========================================================
            # 3) REFRESH UI
            # ==========================================================
            self.view_df = self.df.copy()
            self._populate_detects()
            self._refresh_table()
            self._update_stats_and_chart()
            self._cache_df()
            self._save_cfg()

            self._log(f"Đã nạp thêm {len(paths)} file, tổng {len(self.df)} dòng.")

        except Exception as e:
            messagebox.showerror("Lỗi nạp", str(e))




    def _populate_detects(self):
        if self.df.empty:
            return

        # === Tự động phát hiện cột SO SÁNH và U thực tế ===
        self.compare_col = detect_compare_column(self.df)
        self.voltage_col = pick_voltage_col(self.df)

        # === Chỉ cho phép 2 cột dùng để vẽ: U thực tế và SO SÁNH (%) ===
        valid_plot_cols = []
        if self.compare_col and self.compare_col in self.df.columns:
            valid_plot_cols.append(self.compare_col)
        if self.voltage_col and self.voltage_col in self.df.columns and self.voltage_col not in valid_plot_cols:
            valid_plot_cols.append(self.voltage_col)

        # Cập nhật vào combo box cột vẽ
        self.vcol_cmb.configure(values=valid_plot_cols or [""])
        cur_v = self.vcol_cmb.get().strip()
        if cur_v in valid_plot_cols:
            self.vcol_cmb.set(cur_v)
        elif self.voltage_col in valid_plot_cols:
            self.vcol_cmb.set(self.voltage_col)
        elif valid_plot_cols:
            self.vcol_cmb.set(valid_plot_cols[0])

        # === Danh định ===
        self.nominal_col = pick_nominal_col(self.df)
        if self.nominal_col:
            vals = sorted(self.df[self.nominal_col].dropna().astype(str).unique().tolist())
            self.unom_val_cmb.configure(values=vals if vals else [""])
            cur = self.unom_val_cmb.get().strip()
            if cur and cur in vals:
                self.unom_val_cmb.set(cur)
            elif vals:
                self.unom_val_cmb.set(vals[0])
        else:
            self.unom_val_cmb.configure(values=[""])
            self.unom_val_cmb.set("")

        # === Cột thời gian ===
        self.dt_col = detect_datetime_column(self.df)

        # === Zone_Bx multi-select list ===
        if "Zone_Bx" in self.df.columns:
            self.zones_all = sorted(self.df["Zone_Bx"].dropna().unique().tolist())
        else:
            self.zones_all = []
        # nếu chưa chọn zone nào thì mặc định = tất cả
        if not hasattr(self, "zone_selected"):
            self.zone_selected = set()
        self._update_zone_badge()



    def _apply_filters(self):
        if self.df.empty: return
        df = self.df.copy()
        vcol = self.vcol_cmb.get().strip()
        self.voltage_col = vcol or self.voltage_col

        # station
        text = _norm_text(self.station_text.get())
        station_col = detect_station_column(df)
        if text and station_col:
            col_norm = df[station_col].astype(str).map(_norm_text)
            df = df[col_norm.str.contains(re.escape(text), na=False)].copy()

        # Uđd value filter (exact) if enabled
        if self.use_unom_filter.get():
            nom_col = self.nominal_col
            nom_val = self.unom_val_cmb.get().strip()
            if nom_col and nom_col in df.columns and nom_val:
                df = df[df[nom_col].astype(str) == nom_val].copy()

        # time filter if enabled (inclusive end day)
        if self.use_time_filter.get():
            dt_col = self.dt_col or detect_datetime_column(df)
            if dt_col:
                start = pd.to_datetime(self.from_entry.get_date())
                end   = pd.to_datetime(self.to_entry.get_date())
                dt = pd.to_datetime(df[dt_col], errors="coerce", dayfirst=True)
                df = df.assign(__dt=dt).dropna(subset=["__dt"])
                if pd.notna(start): df = df[df["__dt"] >= start]
                if pd.notna(end):   df = df[df["__dt"] <= (end + pd.Timedelta(days=1) - pd.Timedelta(seconds=1))]
                df = df.drop(columns=["__dt"], errors="ignore")

        # low/high relative filters
        # --- Lọc U THẤP / U CAO theo GIÁ TRỊ U THỰC TẾ ---
        # --- Lọc U THẤP / U CAO so với CỘT SO SÁNH (%) (độc lập với cột U dùng vẽ) ---
        low_on, high_on = self.use_low_filter.get(), self.use_high_filter.get()
        comp_col = self.compare_col or detect_compare_column(df)

        if (low_on or high_on) and comp_col and comp_col in df.columns:
            cmp_series = pd.to_numeric(df[comp_col], errors="coerce")
            df = df.assign(__cmp=cmp_series).dropna(subset=["__cmp"])

            def _to_float(s, default=None):
                try:
                    return float(str(s).replace(",", "."))
                except Exception:
                    return default

            low_thr  = _to_float(self.low_pct_str.get(),  None)  # ví dụ 95
            high_thr = _to_float(self.high_pct_str.get(), None)  # ví dụ 110

            if low_on and low_thr is not None:
                df = df[df["__cmp"] <= low_thr]
            if high_on and high_thr is not None:
                df = df[df["__cmp"] >= high_thr]

            df = df.drop(columns=["__cmp"], errors="ignore")
        elif (low_on or high_on):
            self._log("⚠️ Không tìm thấy cột so sánh (ví dụ 'SO SÁNH (%)'). Vui lòng kiểm tra dữ liệu.")

        # Lọc theo vùng Zone_Bx (multi-select)
        zones = list(getattr(self, "zone_selected", set()) or [])
        if zones and "Zone_Bx" in df.columns:
            df = df[df["Zone_Bx"].isin(zones)]


        # renumber so tt
        if "so tt" in df.columns: df = df.drop(columns=["so tt"], errors="ignore")
        df.insert(0, "so tt", np.arange(1, len(df)+1))

        self.view_df = df
        if comp_col:
            self._log(f"Đang lọc theo cột so sánh: {comp_col}")
        self._refresh_table()
        self._update_stats_and_chart()
        # nếu đang đứng ở Heatmap/Phân phối thì render lại luôn
        try:
            cur_tab = self.dashboard_tabs.get()
            if cur_tab == "Heatmap":
                self._render_heatmap_on_gui()
            elif cur_tab == "Phân phối":
                self._render_dist_on_gui()
            elif cur_tab == "Báo cáo":
                self._render_report_zone_charts_on_gui()
        except Exception:
            pass

        self._cache_df(); self._save_cfg()

    def _maybe_apply_filters(self, *_):
        """Chỉ apply khi Auto đang bật."""
        try:
            if self.auto_apply.get():
                self._apply_filters()
        except Exception:
            pass


    def _open_zone_multiselect(self):
        """Popup chọn nhiều Zone_Bx + search."""
        import tkinter as tk

        # Nếu chưa có danh sách zones thì thôi
        zones = list(getattr(self, "zones_all", []) or [])
        if not zones:
            self._log("⚠️ Chưa có danh sách Zone_Bx để chọn (nạp file trước).")
            return

        win = ctk.CTkToplevel(self)
        win.title("Chọn Zone_Bx")
        win.geometry("420x520")
        win.grab_set()

        # search
        search_var = ctk.StringVar(value="")
        ctk.CTkLabel(win, text="Tìm Zone:", font=("Segoe UI", 13, "bold")).pack(anchor="w", padx=12, pady=(12, 6))
        ent = ctk.CTkEntry(win, textvariable=search_var, placeholder_text="Gõ để lọc…")
        ent.pack(fill="x", padx=12)

        # listbox multi-select (tk widget cho ổn định)
        frame = ctk.CTkFrame(win, fg_color="transparent")
        frame.pack(fill="both", expand=True, padx=12, pady=12)

        lb = tk.Listbox(frame, selectmode=tk.MULTIPLE, activestyle="none")
        lb.pack(side="left", fill="both", expand=True)

        sb = tk.Scrollbar(frame, orient="vertical", command=lb.yview)
        sb.pack(side="right", fill="y")
        lb.configure(yscrollcommand=sb.set)

        # helpers
        def _filtered_list():
            s = search_var.get().strip().lower()
            if not s:
                return zones
            return [z for z in zones if s in str(z).lower()]

        def _render():
            lb.delete(0, tk.END)
            cur = self.zone_selected
            for z in _filtered_list():
                lb.insert(tk.END, z)
            # restore selections in current filtered view
            for i, z in enumerate(_filtered_list()):
                if z in cur:
                    lb.selection_set(i)

        def _select_all():
            self.zone_selected = set(zones)
            _render()

        def _clear_all():
            self.zone_selected = set()
            _render()

        def _apply():
            # lấy selection theo filtered view
            flt = _filtered_list()
            sel_idx = set(lb.curselection())
            picked = {flt[i] for i in sel_idx} if sel_idx else set()

            # nếu đang search, cập nhật subset trong flt; còn lại giữ nguyên
            s = search_var.get().strip()
            if s:
                remain = set(self.zone_selected) - set(flt)
                self.zone_selected = remain | picked
            else:
                self.zone_selected = picked

            # update label hiển thị
            self._update_zone_badge()

            # apply nếu auto bật
            self._maybe_apply_filters()
            win.destroy()

        # buttons
        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(fill="x", padx=12, pady=(0, 12))

        ctk.CTkButton(btn_row, text="Chọn tất cả", command=_select_all, width=110).pack(side="left")
        ctk.CTkButton(btn_row, text="Bỏ chọn", command=_clear_all, width=90).pack(side="left", padx=8)
        ctk.CTkButton(btn_row, text="Áp dụng", command=_apply, width=90).pack(side="right")

        # events
        def _on_search(*_):
            _render()

        search_var.trace_add("write", _on_search)

        _render()
        ent.focus_set()

    def _update_zone_badge(self):
        n = len(getattr(self, "zone_selected", set()) or set())
        if hasattr(self, "zone_badge_lbl"):
            self.zone_badge_lbl.configure(text=(f"{n} zone" if n else "Tất cả"))

    def _display_df(self, df: pd.DataFrame) -> pd.DataFrame:
        return df.drop(columns=[c for c in ["_source_file","_sheet"] if c in df.columns], errors="ignore")

    def _refresh_table(self):
        df_disp = self._display_df(self.view_df.head(5000))
        # ===== FORMAT CỘT NGÀY: dd-mm-yyyy =====
        if "NGÀY" in df_disp.columns:
            try:
                df_disp["NGÀY"] = (
                    pd.to_datetime(df_disp["NGÀY"], errors="coerce", dayfirst=True)
                    .dt.strftime("%d-%m-%Y")
                )
            except Exception:
                pass

        self.table.delete(*self.table.get_children())

        # Thiết lập cột
        self.table["columns"] = list(df_disp.columns)

        self.table.tag_configure("zone_missing", background="#ffe6e6", foreground="red")

        for c in df_disp.columns:
            self.table.heading(c, text=c)
            self.table.column(c, width=90, stretch=True)

        for _, row in df_disp.iterrows():
            values = [str(x) for x in row.tolist()]
            zone_val = row.get("Zone_Bx", None)

            if pd.isna(zone_val):
                self.table.insert("", "end", values=values, tags=("zone_missing",))
            else:
                self.table.insert("", "end", values=values)

        # enable sort by clicking heading
        try:
            self._table_bind_heading_sort()
            self._autofit_table_columns()
        except Exception:
            pass

        # apply zebra rows
        for idx, iid in enumerate(self.table.get_children("")):
            self.table.item(iid, tags=("even" if idx % 2 == 0 else "odd"))

    def _autofit_table_columns(self, max_width=420, min_width=60, padding=14):
        """
        Auto-fit column width cho ttk.Treeview dựa trên:
        - độ dài tiêu đề cột
        - độ dài nội dung các ô
        Giới hạn min/max để tránh cột quá to.
        """
        import tkinter.font as tkfont

        if not hasattr(self, "table"):
            return

        tree = self.table
        font = tkfont.Font(font=("Segoe UI", 10))

        for col in tree["columns"]:
            # độ rộng theo heading
            header_text = col
            width = font.measure(header_text) + padding

            # độ rộng theo nội dung (lấy sample để nhanh)
            for iid in tree.get_children("")[:200]:  # limit 200 rows
                cell = tree.set(iid, col)
                if cell is None:
                    continue
                w = font.measure(str(cell)) + padding
                if w > width:
                    width = w

            width = max(min_width, min(width, max_width))
            tree.column(col, width=width, stretch=False)

    def _export_missing_tba(self):
        if self.df.empty or "Zone_Bx" not in self.df.columns:
            messagebox.showwarning("Thiếu dữ liệu", "Chưa có dữ liệu hoặc chưa gắn được Zone_Bx.")
            return

        # Lọc các dòng thiếu Zone_Bx và lấy danh sách TBA duy nhất
        tba_col = detect_station_column(self.df)
        if not tba_col:
            messagebox.showwarning("Thiếu cột", "Không tìm thấy cột TRẠM BIẾN ÁP.")
            return

        df_missing = self.df[self.df["Zone_Bx"].isna()]
        tba_missing = df_missing[tba_col].dropna().astype(str).str.strip().unique()

        if len(tba_missing) == 0:
            messagebox.showinfo("OK", "Tất cả trạm đã được ánh xạ Zone_Bx.")
            return

        # Chọn nơi lưu file
        save_path = filedialog.asksaveasfilename(
            title="Lưu danh sách TBA lỗi",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")]
        )
        if not save_path:
            return

        try:
            df_out = pd.DataFrame({"TRẠM BIẾN ÁP lỗi": sorted(tba_missing)})
            df_out.to_excel(save_path, index=False)
            messagebox.showinfo("[OK] Đã lưu", f"Đã lưu danh sách {len(df_out)} TBA lỗi vào:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Lỗi khi xuất", str(e))
    def _update_kpi_cards(self):
        """Cập nhật KPI cards + dòng thống kê dựa trên view_df hiện tại."""
        df = self.view_df if hasattr(self, "view_df") and not self.view_df.empty else self.df

        so_tba = 0
        umin = utb = umax = np.nan

        try:
            if df is not None and not df.empty:
                station_col = detect_station_column(df)
                if station_col and station_col in df.columns:
                    so_tba = int(df[station_col].nunique())

                # ưu tiên cột đang chọn để vẽ; fallback tự dò
                vcol = None
                vsel = ""
                try:
                    vsel = self.vcol_cmb.get().strip()
                except Exception:
                    vsel = ""
                if vsel and vsel in df.columns:
                    vcol = vsel
                elif getattr(self, "voltage_col", None) in df.columns:
                    vcol = self.voltage_col
                else:
                    vcol = pick_voltage_col(df)

                if vcol and vcol in df.columns:
                    v = pd.to_numeric(df[vcol], errors="coerce").dropna()
                    if not v.empty:
                        umin, utb, umax = float(v.min()), float(v.mean()), float(v.max())
        except Exception as e:
            safe_print("KPI update error:", e)

        # update vars
        if hasattr(self, "kpi_vars"):
            self.kpi_vars["rows"].set(f"{0 if df is None else int(len(df))}")
            self.kpi_vars["tba"].set(f"{so_tba:d}")

            def _fmt2(x):
                return "0" if pd.isna(x) else f"{x:.2f}"

            self.kpi_vars["umin"].set(_fmt2(umin))
            self.kpi_vars["utb"].set(_fmt2(utb))
            self.kpi_vars["umax"].set(_fmt2(umax))

        if hasattr(self, "stats_var"):
            n = 0 if df is None else int(len(df))

            def _fmt1(x):
                return "—" if pd.isna(x) else f"{x:.1f}"

            self.stats_var.set(f"Thống kê: {n} dòng | Umin={_fmt1(umin)}  Utb={_fmt1(utb)}  Umax={_fmt1(umax)}")
    def _update_stats_and_chart(self):
        # luôn cập nhật KPI theo view_df hiện tại
        self._update_kpi_cards()

        vcol = self.vcol_cmb.get().strip() if hasattr(self, "vcol_cmb") else ""
        if not vcol:
            vcol = self.voltage_col

        if self.view_df.empty or not vcol or vcol not in self.view_df.columns:
            self._draw_chart_empty()
            return

        v = pd.to_numeric(self.view_df[vcol], errors="coerce").dropna()
        if v.empty:
            self._draw_chart_empty()
            return

        self._draw_chart()

    def _draw_chart_empty(self):
        self.ax.cla()
        vcol = self.vcol_cmb.get().strip() if hasattr(self, "vcol_cmb") else ""
        if not vcol:
            vcol = self.voltage_col
        title = f"Biểu đồ {vcol}" if vcol else "Biểu đồ"
        self.ax.set_title(title)
        self.ax.set_xlabel("Thời gian / Index")
        self.ax.set_ylabel("Điện áp")
        self.canvas.draw()


    def _draw_chart(self):
        import matplotlib.dates as mdates

        self.ax.cla()
        vcol = self.vcol_cmb.get().strip() or self.voltage_col

        self.ax.set_title(f"Biểu đồ {vcol}" if vcol else "Biểu đồ")
        self.ax.set_ylabel("Điện áp")

        data = self.view_df.copy()
        if vcol not in data.columns:
            safe_print("[x] Không tìm thấy cột U THỰC TẾ trong dữ liệu.")
            self.canvas.draw()
            return

        # Làm sạch dữ liệu
        data["__v"] = pd.to_numeric(data[vcol], errors="coerce")
        data = data.dropna(subset=["__v"])

        safe_print("[[OK]] Số điểm hợp lệ để vẽ:", len(data))
        if data.empty:
            self.canvas.draw()
            return

        # Xử lý cột thời gian
        dt_col = detect_datetime_column(data)
        if dt_col:
            data["__x"] = pd.to_datetime(data[dt_col], errors="coerce", dayfirst=True)
            data = data.dropna(subset=["__x"]).sort_values("__x")
            xvals = data["__x"].values
            self.ax.set_xlabel(f"Thời gian ({dt_col})")

            # Format thời gian đẹp
            self.ax.xaxis.set_major_locator(mdates.AutoDateLocator())
            span_days = (data["__x"].max() - data["__x"].min()).days
            if span_days <= 2:
                self.ax.xaxis.set_major_formatter(mdates.DateFormatter("%d-%m %H:%M"))
            else:
                self.ax.xaxis.set_major_formatter(mdates.DateFormatter("%d-%m"))
            self.fig.autofmt_xdate(rotation=45)

        else:
            data = data.reset_index(drop=True)
            xvals = data.index.values
            self.ax.set_xlabel("Index")
            safe_print("[⚠️] Không có cột thời gian — dùng index thay x.")

        # Vẽ scatter hoặc line
        if self.chart_mode.get() == "scatter":
            self.ax.scatter(xvals, data["__v"].values, s=8, alpha=0.7)
        else:
            self.ax.plot(xvals, data["__v"].values, lw=1)

        self.ax.grid(True, linestyle="--", alpha=0.3)
        self.canvas.draw()


    def _export_figure(self):
        if self.view_df.empty:
            messagebox.showwarning("Rỗng","Không có dữ liệu để xuất hình."); return
        initial = self.last_dir if os.path.isdir(self.last_dir) else os.path.expanduser("~")
        out = filedialog.asksaveasfilename(title="Lưu hình PNG", initialdir=initial,
                                           defaultextension=".png", filetypes=[("PNG Image","*.png")])
        if not out: return
        try:
            self.fig.savefig(out, dpi=160, bbox_inches="tight"); self._log(f"Đã lưu hình: {out}")
            self.last_dir = os.path.dirname(out)
        except Exception as e:
            messagebox.showerror("Lỗi lưu", str(e))

    # ---------- EXPORTS ----------


    # ---- draw/update helpers ----
    def _on_close(self):
        self._save_cfg(); self._cache_df(); self.destroy()
    def detect_compare_column(df: pd.DataFrame) -> Optional[str]:
        for c in df.columns:
            low = str(c).lower()
            if ("so sánh" in low) or ("so sanh" in low) or ("%" in low):
                if pd.api.types.is_numeric_dtype(df[c]):
                    return c
        return None

    def _plot_voltage_heatmap(self):
        if self.view_df.empty:
            messagebox.showwarning("Thiếu dữ liệu", "Không có dữ liệu để vẽ heatmap.")
            return

        df = self.view_df.copy()
        dt_col = self.dt_col or detect_datetime_column(df)
        vcol = self.voltage_col

        if not dt_col or dt_col not in df.columns or not vcol or vcol not in df.columns:
            messagebox.showwarning("Thiếu cột", "Chưa xác định được cột thời gian hoặc điện áp.")
            return

        # Ghép ngày + giờ nếu có cột Giờ riêng
        hour_col = None
        for c in df.columns:
            if "giờ" in c.lower() or "hour" in c.lower():
                hour_col = c
                break

        dt = pd.to_datetime(df[dt_col], errors="coerce", dayfirst=True)
        if hour_col:
            hour_val = pd.to_numeric(df[hour_col], errors="coerce").fillna(0)
            dt += pd.to_timedelta(hour_val, unit="h")

        v = pd.to_numeric(df[vcol], errors="coerce")
        df["__date"] = dt.dt.date
        df["__hour"] = dt.dt.hour
        df["__v"] = v
        tmp = df.dropna(subset=["__date", "__hour", "__v"])

        if tmp.empty:
            messagebox.showwarning("Dữ liệu trống", "Không có giá trị hợp lệ để vẽ.")
            return

        pivot = tmp.pivot_table(index="__hour", columns="__date", values="__v", aggfunc="mean")
        pivot = pivot.reindex(range(24))  # đảm bảo đủ 0–23h

        import matplotlib.pyplot as plt
        import seaborn as sns

        plt.figure(figsize=(12, 6))
        ax = sns.heatmap(pivot, cmap="YlGnBu", cbar_kws={"label": "U thực tế (kV)"})
        ax.set_title("Heatmap U thực tế trung bình theo Giờ và Ngày")
        ax.set_xlabel("Ngày")
        ax.set_ylabel("Giờ")
        plt.tight_layout()
        plt.show()


    def _plot_voltage_hist_box(self):
        if self.view_df.empty:
            messagebox.showwarning("Thiếu dữ liệu", "Không có dữ liệu để vẽ.")
            return

        vcol = self.voltage_col
        if not vcol or vcol not in self.view_df.columns:
            messagebox.showwarning("Thiếu cột", "Chưa xác định được cột U thực tế.")
            return

        v = pd.to_numeric(self.view_df[vcol], errors="coerce").dropna()
        if v.empty:
            messagebox.showwarning("Dữ liệu rỗng", "Không có giá trị điện áp hợp lệ.")
            return

        import matplotlib.pyplot as plt

        fig, axs = plt.subplots(1, 2, figsize=(12, 5))
        fig.suptitle(f"Phân tích phân phối U thực tế ({vcol})", fontsize=14)

        axs[0].hist(v, bins=30, color="skyblue", edgecolor="black")
        axs[0].set_title("Histogram U thực tế")
        axs[0].set_xlabel("U (kV)")
        axs[0].set_ylabel("Số lần")

        axs[1].boxplot(v, vert=True, patch_artist=True, boxprops=dict(facecolor="lightgreen"))
        axs[1].set_title("Boxplot U thực tế")
        axs[1].set_ylabel("U (kV)")

        plt.tight_layout()
        plt.show()

    def _show_dashboard_fix_tba_loi(self):
        import pandas as pd
        import webview
        import os
        import tempfile
        from rapidfuzz import process, fuzz

        #db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "DB_VietSub.xlsx")
        db_path = get_db_path()

        db_buses = pd.read_excel(db_path, sheet_name="Buses")
        tba_scada_set = set(db_buses["TBA_SCADA"].astype(str).str.strip().str.lower())

        df = self.view_df.copy()
        tba_col = detect_station_column(df)
        tba_all = df[tba_col].astype(str).unique()
        tba_loi = [tba for tba in tba_all if tba.strip().lower() not in tba_scada_set]
        if not tba_loi:
            from tkinter import messagebox
            messagebox.showinfo("OK", "Không còn TBA lỗi nào! Bạn có thể xem báo cáo tổng hợp.")
            return

        html = """
        <html><head>
            <meta charset="utf-8">
            <style>
                body { font-family: Arial; background: #fafdff; }
                table { border-collapse: collapse; font-size: 15px; }
                th,td { border: 1px solid #bbb; padding: 6px 10px; }
                th { background: #e8e8e8; }
                .tba-loi { background: #FFF98C }
                .highlight { background: #c4ffa3 !important; }
                .dash-btn { margin: 10px 0; padding: 7px 15px; background: #1756d9; color: #fff; border-radius: 6px;}
            </style>
        </head>
        <body>
            <h2 style="color:#d95f05;">DASHBOARD HIỆU CHỈNH TBA LỖI</h2>
            <table>
                <tr><th>STT</th><th>TBA Lỗi</th><th>Gợi ý tên đúng (chọn 1 để sửa)</th><th>Sửa</th></tr>
        """
        for idx, tba in enumerate(tba_loi, 1):
            suggests = process.extract(tba, db_buses["TBA_SCADA"].astype(str).tolist(), limit=5, scorer=fuzz.ratio)
            suggest_html = ""
            group_name = f"tba_suggest_{idx}"
            for s in suggests:
                pct = f"{s[1]:.1f}"
                suggest_html += f"<label><input type='radio' name='{group_name}' value='{s[0]}'> {s[0]} ({pct}%)</label><br>"
            suggest_html += f"<label><input type='radio' name='{group_name}' value=''> (Không có trong DB)</label>"
            html += f"""<tr>
                <td>{idx}</td>
                <td class='tba-loi'>{tba}</td>
                <td>{suggest_html}</td>
                <td><button onclick="submitEditTBA('{tba}','{group_name}',this)">Sửa</button></td>
            </tr>"""
        html += """
            </table>
            <br><button class="dash-btn" onclick="window.location.reload()">Làm mới danh sách</button>
            <script>
            function submitEditTBA(tba, group, btn){
                let radios = document.getElementsByName(group);
                let new_tba = "";
                for(let i=0;i<radios.length;i++) if(radios[i].checked) new_tba = radios[i].value;
                if(!new_tba){
                    alert("Chọn 1 tên đúng để sửa (hoặc tự cập nhật trong DB nếu không có)");
                    return;
                }
                window.pywebview.api.update_tba_scada(tba, new_tba).then(function(msg){
                    btn.outerHTML = "<span style='color:#0a0; font-weight:bold;'>✓ Đã cập nhật!</span>";
                    btn.closest("tr").classList.add("highlight");
                    alert(msg);
                });
            }
            </script>
        </body></html>
        """

        class Api:
            def update_tba_scada(self, old_tba, new_tba):
                import os, shutil, tempfile, datetime
                import openpyxl
                from openpyxl.styles import PatternFill

                def _acquire_lock(lock_path: str):
                    # lock file đơn giản, tránh ghi trùng
                    fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
                    os.write(fd, b"lock")
                    os.close(fd)

                def _release_lock(lock_path: str):
                    try:
                        if os.path.exists(lock_path):
                            os.remove(lock_path)
                    except Exception:
                        pass

                def _backup_db(db_path: str) -> str:
                    app_dir = os.path.dirname(os.path.abspath(db_path))
                    backup_dir = os.path.join(app_dir, "DB_backups")
                    os.makedirs(backup_dir, exist_ok=True)
                    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    base = os.path.splitext(os.path.basename(db_path))[0]
                    backup_path = os.path.join(backup_dir, f"{base}_backup_{ts}.xlsx")
                    shutil.copy2(db_path, backup_path)
                    return backup_path

                def _atomic_save_workbook(wb, db_path: str):
                    # save ra file tạm cùng thư mục rồi replace -> an toàn hơn
                    folder = os.path.dirname(os.path.abspath(db_path))
                    fd, tmp_path = tempfile.mkstemp(prefix="~tmp_db_", suffix=".xlsx", dir=folder)
                    os.close(fd)
                    try:
                        wb.save(tmp_path)
                        os.replace(tmp_path, db_path)
                    finally:
                        try:
                            if os.path.exists(tmp_path):
                                os.remove(tmp_path)
                        except Exception:
                            pass

                old_k = str(old_tba).strip()
                new_k = str(new_tba).strip()
                if not old_k or not new_k:
                    return "Thiếu dữ liệu old/new."

                lock_path = db_path + ".lock"

                try:
                    _acquire_lock(lock_path)
                except FileExistsError:
                    return "DB đang được chỉnh sửa ở nơi khác. Hãy đóng các cửa sổ/tool khác rồi thử lại."

                try:
                    # 1) backup trước khi đụng DB
                    backup_path = _backup_db(db_path)

                    # 2) load + kiểm tra cấu trúc tối thiểu
                    wb = openpyxl.load_workbook(db_path)
                    if "Buses" not in wb.sheetnames:
                        return "Không tìm thấy sheet 'Buses' trong DB_VietSub."

                    ws = wb["Buses"]

                    # tìm cột TBA_SCADA đúng theo header hàng 1
                    col_scada = None
                    for i, cell in enumerate(ws[1], start=1):
                        if str(cell.value).strip() == "TBA_SCADA":
                            col_scada = i
                            break
                    if not col_scada:
                        return "Không tìm thấy cột TBA_SCADA trong sheet Buses."

                    # 3) update giá trị: tìm những dòng có TBA_SCADA == new_k => đổi về old_k (theo logic tool đang dùng)
                    updated = 0
                    for r in range(2, ws.max_row + 1):
                        v = ws.cell(row=r, column=col_scada).value
                        if v is None:
                            continue
                        if str(v).strip().lower() == new_k.lower():
                            ws.cell(row=r, column=col_scada).value = old_k
                            ws.cell(row=r, column=col_scada).fill = PatternFill("solid", fgColor="FFF200")
                            updated += 1

                    if updated == 0:
                        return f"Không tìm thấy '{new_k}' trong cột TBA_SCADA để đổi sang '{old_k}'. (Không ghi DB)"

                    # 4) tuyệt đối KHÔNG thêm cột mới (tránh phá DB)
                    # -> bỏ hẳn ws.max_column+1

                    # 5) ghi DB kiểu atomic
                    _atomic_save_workbook(wb, db_path)

                    return f"✅ Đã cập nhật {updated} dòng: {new_k} ➜ {old_k}. (Backup: {os.path.basename(backup_path)})"

                except Exception as e:
                    return f"❌ Lỗi cập nhật DB (đã có backup): {e}"
                finally:
                    _release_lock(lock_path)


        api = Api()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".html", mode="w", encoding="utf-8") as f:
            f.write(html)
            html_path = f.name

        webview.create_window("Hiệu chỉnh TBA lỗi", html_path, width=820, height=700, js_api=api)
        webview.start()

        if os.path.exists(html_path):
            os.remove(html_path)




    def _show_dashboard_zone_voltage_report(self):
        import plotly.graph_objs as go
        import plotly.io as pio
        import webview
        import tempfile
        import os
        import pandas as pd
        self.status_var.set("⏳ Đang tạo báo cáo Dashboard...")
        self.update_idletasks()

        # ===== Thêm check TBA lỗi theo DB VietSub =====
        #db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "DB_VietSub.xlsx")
        db_path = get_db_path()

        db_buses = pd.read_excel(db_path, sheet_name="Buses")
        tba_scada_set = set(db_buses["TBA_SCADA"].astype(str).str.strip().str.lower())
        def is_tba_loi(tba_name):
            return tba_name.strip().lower() not in tba_scada_set

        df = self.view_df.copy()
        if df.empty:
            from tkinter import messagebox
            messagebox.showwarning("Rỗng", "Không có dữ liệu để hiển thị."); return

        vcol = self.voltage_col
        un_col = self.nominal_col if self.nominal_col in df.columns else None
        station_col = detect_station_column(df)
        if not vcol or not un_col or not station_col:
            from tkinter import messagebox
            messagebox.showwarning("Thiếu cột", "Không tìm thấy cột cần thiết."); return

        # === Lấy thời gian lọc từ widget DateEntry (from_entry, to_entry) ===
        from_date = self.from_entry.get_date()
        to_date   = self.to_entry.get_date()
        def _format_date(d):
            try:
                return pd.to_datetime(d).strftime("%d/%m/%Y")
            except Exception:
                return str(d)
        if from_date and to_date:
            time_label = f"{_format_date(from_date)} - {_format_date(to_date)}"
            file_time = f"{pd.to_datetime(from_date).strftime('%Y-%m-%d')}_{pd.to_datetime(to_date).strftime('%Y-%m-%d')}"
        else:
            time_label = ""
            file_time = pd.Timestamp.today().strftime("%Y-%m-%d")

        df["Ut"] = pd.to_numeric(df[vcol], errors="coerce")
        df["Un"] = pd.to_numeric(df[un_col], errors="coerce")
        df = df.dropna(subset=["Ut","Un"])
        if "Zone_Bx" not in df.columns:
            df["Zone_Bx"] = "(Chưa có Zone)"
        zone_col = "Zone_Bx"


        # ----- Tổng hợp bảng chi tiết CAO & THẤP -----
        high_rows, low_rows = [], []
        for zone, df_zone in df.groupby(zone_col):
            for (tba, udinh), df_tba_udinh in df_zone.groupby([station_col, "Un"]):
                ut = df_tba_udinh["Ut"]
                over_mask = ut >= 1.10 * udinh
                if sum(over_mask) > 0:
                    umax_cao = float(ut[over_mask].max())
                    row = {
                        "STT": None,
                        "Zone_Bx": zone,
                        "TBA": tba,
                        "Udđ": udinh,
                        "U(kV)": umax_cao,
                        "U(kV)/Udđ": round(umax_cao / udinh, 3),
                        "Số lần Cao": int(sum(over_mask)),
                        "Umin": float(ut[over_mask].min()),
                        "Umax": umax_cao
                    }
                    high_rows.append(row)
                under_mask = ut <= 0.95 * udinh
                if sum(under_mask) > 0:
                    umin_thap = float(ut[under_mask].min())
                    row = {
                        "STT": None,
                        "Zone_Bx": zone,
                        "TBA": tba,
                        "Udđ": udinh,
                        "U(kV)": umin_thap,
                        "U(kV)/Udđ": round(umin_thap / udinh, 3),
                        "Số lần Thấp": int(sum(under_mask)),
                        "Umin": umin_thap,
                        "Umax": float(ut[under_mask].max())
                    }
                    low_rows.append(row)

        df_high = pd.DataFrame(high_rows)
        df_low = pd.DataFrame(low_rows)
        if not df_high.empty:
            df_high = df_high.sort_values(["Zone_Bx", "Số lần Cao"], ascending=[True, False]).reset_index(drop=True)
            df_high["STT"] = range(1, len(df_high)+1)
        if not df_low.empty:
            df_low = df_low.sort_values(["Zone_Bx", "Số lần Thấp"], ascending=[True, False]).reset_index(drop=True)
            df_low["STT"] = range(1, len(df_low)+1)

        # --- DANH SÁCH TBA LỖI ở cả CAO & THẤP ---
        tba_loi_high = list(df_high[df_high["TBA"].apply(is_tba_loi)]["TBA"].unique())
        tba_loi_low  = list(df_low[df_low["TBA"].apply(is_tba_loi)]["TBA"].unique())
        tba_loi_set = set(tba_loi_high) | set(tba_loi_low)

        # ==== Cảnh báo TBA lỗi (nếu còn) ====
        if tba_loi_set:
            tba_loi_html = "<div style='padding:10px; border:2px solid #F05; background:#FFF6E6; color:#F05; border-radius:10px; margin-bottom:18px;'>"
            tba_loi_html += "<b>⚠️ DANH SÁCH TBA ĐANG LỖI TÊN (chưa được tổng hợp):</b><br>"
            tba_loi_html += "<ul style='margin:8px 0 0 20px;'>"
            for tba in sorted(tba_loi_set):
                tba_loi_html += f"<li style='margin-bottom:3px;'><b>{tba}</b></li>"
            tba_loi_html += "</ul>"
            tba_loi_html += "<div style='margin-top:6px; color:#888; font-size:14px;'>Hãy sửa tên TBA này ở dashboard hiệu chỉnh để báo cáo tổng hợp đủ!</div>"
            tba_loi_html += "</div>"
        else:
            tba_loi_html = ""

        # --- Tổng hợp Zone_Bx cho CAO & THẤP ---
        def zone_stat(df, num_col, label_sum, label_tba):
            if df.empty:
                return pd.DataFrame(columns=["Zone_Bx", label_tba, label_sum])
            return df.groupby("Zone_Bx").agg(
                **{label_tba: ("TBA", "nunique"),
                   label_sum: (num_col, "sum")}
            ).reset_index()

        stat_high = zone_stat(df_high, "Số lần Cao", "Tổng số lần Cao", "Số TBA vi phạm")
        stat_low  = zone_stat(df_low, "Số lần Thấp", "Tổng số lần Thấp", "Số TBA vi phạm")

        # --- Biểu đồ Zone_Bx CAO/THẤP ---
        fig_high1 = go.Figure()
        fig_high1.add_trace(go.Bar(
            x=stat_high["Zone_Bx"], y=stat_high["Số TBA vi phạm"],
            marker_color="rgb(36,190,110)",
            text=stat_high["Số TBA vi phạm"], textposition="auto",
            name="Số TBA vi phạm CAO",
            hovertemplate='<b>%{x}</b><br>Số TBA vi phạm: %{y}'
        ))
        fig_high1.update_layout(
            title="Số lượng TBA có vi phạm điện áp CAO (>=110%) theo Zone_Bx",
            xaxis_title="Zone_Bx", yaxis_title="Số TBA vi phạm",
            height=320, font=dict(family="Arial", size=13), plot_bgcolor="#fafdff"
        )
        fig_high2 = go.Figure()
        fig_high2.add_trace(go.Bar(
            x=stat_high["Zone_Bx"], y=stat_high["Tổng số lần Cao"],
            marker_color="rgb(0,120,250)",
            text=stat_high["Tổng số lần Cao"], textposition="auto",
            name="Tổng số lần vi phạm CAO",
            hovertemplate='<b>%{x}</b><br>Tổng số lần Cao: %{y}'
        ))
        fig_high2.update_layout(
            title="Tổng số lần vi phạm điện áp CAO (>=110%) theo Zone_Bx",
            xaxis_title="Zone_Bx", yaxis_title="Tổng số lần Cao",
            height=320, font=dict(family="Arial", size=13), plot_bgcolor="#fafdff"
        )

        fig_low1 = go.Figure()
        fig_low1.add_trace(go.Bar(
            x=stat_low["Zone_Bx"], y=stat_low["Số TBA vi phạm"],
            marker_color="rgb(241,98,53)",
            text=stat_low["Số TBA vi phạm"], textposition="auto",
            name="Số TBA vi phạm THẤP",
            hovertemplate='<b>%{x}</b><br>Số TBA vi phạm: %{y}'
        ))
        fig_low1.update_layout(
            title="Số lượng TBA có vi phạm điện áp THẤP (<=95%) theo Zone_Bx",
            xaxis_title="Zone_Bx", yaxis_title="Số TBA vi phạm",
            height=320, font=dict(family="Arial", size=13), plot_bgcolor="#fafdff"
        )
        fig_low2 = go.Figure()
        fig_low2.add_trace(go.Bar(
            x=stat_low["Zone_Bx"], y=stat_low["Tổng số lần Thấp"],
            marker_color="rgb(140, 75, 230)",
            text=stat_low["Tổng số lần Thấp"], textposition="auto",
            name="Tổng số lần vi phạm THẤP",
            hovertemplate='<b>%{x}</b><br>Tổng số lần Thấp: %{y}'
        ))
        fig_low2.update_layout(
            title="Tổng số lần vi phạm điện áp THẤP (<=95%) theo Zone_Bx",
            xaxis_title="Zone_Bx", yaxis_title="Tổng số lần Thấp",
            height=320, font=dict(family="Arial", size=13), plot_bgcolor="#fafdff"
        )



        # ======= Bảng HTML đẹp =======
        def table_html(df, caption=""):
            if df.empty:
                return f"<i>Không có số liệu.</i>"
            html = df.to_html(index=False, classes="table table-striped", border=1, float_format="%.3f")
            return (f"<div style='font-weight:bold;margin:8px 0'>{caption}</div>{html}")

        html_high_stat = table_html(stat_high, "BẢNG TK ZONE_BX: Số TBA & số lần VI PHẠM ĐIỆN ÁP CAO (>=110%)")
        html_low_stat  = table_html(stat_low,  "BẢNG TK ZONE_BX: Số TBA & số lần VI PHẠM ĐIỆN ÁP THẤP (<=95%)")
        html_high_detail = table_html(df_high, "BẢNG CHI TIẾT TBA VI PHẠM ĐIỆN ÁP CAO (>=110%)")
        html_low_detail  = table_html(df_low,  "BẢNG CHI TIẾT TBA VI PHẠM ĐIỆN ÁP THẤP (<=95%)")

        fig_high1_html = pio.to_html(fig_high1, full_html=False, include_plotlyjs=True)
        fig_high2_html = pio.to_html(fig_high2, full_html=False, include_plotlyjs=False)
        fig_low1_html  = pio.to_html(fig_low1,  full_html=False, include_plotlyjs=False)
        fig_low2_html  = pio.to_html(fig_low2,  full_html=False, include_plotlyjs=False)

        # ==== HTML dashboard có cảnh báo TBA lỗi ====
        html = f"""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial; margin: 16px; background: #fafdff; }}
                .table {{ border-collapse: collapse; width: 100%; font-size: 15px; margin-bottom:18px;}}
                .table th, .table td {{ border: 1px solid #bbb; padding: 4px 8px; }}
                .table th {{ background: #e8e8e8; }}
                .dash-btn {{
                    display:inline-block; margin:12px 0; padding:7px 15px;
                    background:#1756d9; color:#fff; border:none; border-radius:6px; font-size:15px; cursor:pointer;}}
                h2,h3 {{margin-top:18px;}}
            </style>
        </head>
        <body>
            <h2 style="color:#1756d9">
                BÁO CÁO PHÂN TÍCH ĐIỆN ÁP THEO ZONE_BX
                <br><span style="font-size:18px;color:#F05;">{time_label}</span>
            </h2>
            {tba_loi_html}
            {fig_high1_html}
            {fig_high2_html}
            {html_high_stat}
            <hr>
            {fig_low1_html}
            {fig_low2_html}
            {html_low_stat}
            <hr>
            {html_high_detail}
            <hr>
            {html_low_detail}
            <button class="dash-btn" onclick="window.pywebview.api.export_excel()">📥 Xuất báo cáo Excel</button>
            <button class="dash-btn" onclick="window.pywebview.api.export_word()">📝 Xuất Word báo cáo</button>
        </body>
        </html>
        """

        # --- Lưu biểu đồ ra file PNG để chèn vào Word ---
        chart_imgs = []
        with tempfile.TemporaryDirectory() as tmpdir:
            fig_high1_path = os.path.join(tmpdir, "zone_high1.png")
            fig_high2_path = os.path.join(tmpdir, "zone_high2.png")
            fig_low1_path = os.path.join(tmpdir, "zone_low1.png")
            fig_low2_path = os.path.join(tmpdir, "zone_low2.png")
            fig_high1.write_image(fig_high1_path, scale=2, width=950, height=340)
            fig_high2.write_image(fig_high2_path, scale=2, width=950, height=340)
            fig_low1.write_image(fig_low1_path, scale=2, width=950, height=340)
            fig_low2.write_image(fig_low2_path, scale=2, width=950, height=340)
            chart_imgs = [fig_high1_path, fig_high2_path, fig_low1_path, fig_low2_path]

            with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as f:
                f.write(html.encode("utf-8"))
                html_path = f.name

            class Api:
                def export_excel(self):
                    import os
                    app_dir = os.path.dirname(os.path.abspath(__file__))
                    file_path = os.path.join(app_dir, f"Báo cáo {file_time}.xlsx")
                    with pd.ExcelWriter(file_path, engine="openpyxl") as writer:
                        if not stat_high.empty:
                            stat_high.to_excel(writer, sheet_name="ZONE_BX_HIGH_STAT", index=False)
                        if not stat_low.empty:
                            stat_low.to_excel(writer, sheet_name="ZONE_BX_LOW_STAT", index=False)
                        if not df_high.empty:
                            df_high.to_excel(writer, sheet_name="HIGH_VOLTAGE_DETAIL", index=False)
                        if not df_low.empty:
                            df_low.to_excel(writer, sheet_name="LOW_VOLTAGE_DETAIL", index=False)
                    webview.windows[0].evaluate_js("alert('Đã xuất báo cáo Excel thành công!');")
                    os.startfile(file_path)

                def export_word(self):
                    import os
                    from docx import Document
                    from docx.shared import Inches
                    app_dir = os.path.dirname(os.path.abspath(__file__))
                    file_path = os.path.join(app_dir, f"Báo cáo {file_time}.docx")
                    doc = Document()
                    # Nếu muốn cảnh báo TBA lỗi xuất ra luôn file Word, thêm đoạn này:
                    if tba_loi_html:
                        doc.add_paragraph("⚠️ DANH SÁCH TBA LỖI: " + ", ".join(sorted(tba_loi_set)), style="Intense Quote")
                    doc.add_heading(f'BÁO CÁO PHÂN TÍCH ĐIỆN ÁP THEO ZONE_BX\n{time_label}', 0)
                    doc.add_heading('Biểu đồ tổng hợp điện áp CAO', level=1)
                    doc.add_picture(chart_imgs[0], width=Inches(6.2))
                    doc.add_picture(chart_imgs[1], width=Inches(6.2))
                    doc.add_paragraph()
                    doc.add_heading('Biểu đồ tổng hợp điện áp THẤP', level=1)
                    doc.add_picture(chart_imgs[2], width=Inches(6.2))
                    doc.add_picture(chart_imgs[3], width=Inches(6.2))
                    doc.add_paragraph()

                    doc.add_heading('Thống kê điện áp CAO', level=1)
                    if not stat_high.empty:
                        t = doc.add_table(rows=1, cols=len(stat_high.columns), style='Table Grid')
                        for j, col in enumerate(stat_high.columns):
                            t.cell(0, j).text = str(col)
                        for idx, row in stat_high.iterrows():
                            cells = t.add_row().cells
                            for j, val in enumerate(row):
                                cells[j].text = str(val)
                        doc.add_paragraph()
                    doc.add_heading('Thống kê điện áp THẤP', level=1)
                    if not stat_low.empty:
                        t = doc.add_table(rows=1, cols=len(stat_low.columns), style='Table Grid')
                        for j, col in enumerate(stat_low.columns):
                            t.cell(0, j).text = str(col)
                        for idx, row in stat_low.iterrows():
                            cells = t.add_row().cells
                            for j, val in enumerate(row):
                                cells[j].text = str(val)
                        doc.add_paragraph()
                    doc.add_heading('Bảng chi tiết TBA vi phạm điện áp CAO', level=1)
                    if not df_high.empty:
                        t = doc.add_table(rows=1, cols=len(df_high.columns), style='Table Grid')
                        for j, col in enumerate(df_high.columns):
                            t.cell(0, j).text = str(col)
                        for idx, row in df_high.iterrows():
                            cells = t.add_row().cells
                            for j, val in enumerate(row):
                                cells[j].text = str(val)
                        doc.add_paragraph()
                    doc.add_heading('Bảng chi tiết TBA vi phạm điện áp THẤP', level=1)
                    if not df_low.empty:
                        t = doc.add_table(rows=1, cols=len(df_low.columns), style='Table Grid')
                        for j, col in enumerate(df_low.columns):
                            t.cell(0, j).text = str(col)
                        for idx, row in df_low.iterrows():
                            cells = t.add_row().cells
                            for j, val in enumerate(row):
                                cells[j].text = str(val)
                        doc.add_paragraph()
                    doc.save(file_path)
                    webview.windows[0].evaluate_js("alert('Đã xuất báo cáo Word!');")
                    os.startfile(file_path)

            api = Api()
            webview.create_window("Dashboard & Báo cáo tổng hợp", html_path, width=1300, height=950, js_api=api)
            webview.start()
            if os.path.exists(html_path):
                os.remove(html_path)


    def _show_help(self):
        win = ctk.CTkToplevel(self)
        win.title("Hướng dẫn sử dụng & Bản quyền")
        win.geometry("600x540")
        win.resizable(False, False)

        # Đưa cửa sổ Help lên trước GUI
        win.lift()
        win.attributes("-topmost", True)
        win.after(200, lambda: win.attributes("-topmost", False))  # chỉ giữ trên cùng lúc mở

        # Tiêu đề
        ctk.CTkLabel(
            win, text="📖 HƯỚNG DẪN SỬ DỤNG",
            font=("Segoe UI", 22, "bold"),
            text_color="#1a2857"
        ).pack(pady=(18, 10))

        # Nội dung chi tiết
        help_text = (
            "1. Chức năng chính:\n"
            "   • 📁 Nạp file: Chọn một hoặc nhiều file Excel để phân tích dữ liệu\n"
            "   • 🧹 Xóa: Xóa dữ liệu hiện tại khỏi bảng & biểu đồ\n"
            "   • 🛠️ Hiệu chỉnh TBA lỗi: Mở dashboard web để dò/sửa TBA chưa khớp DB\n"
            "   • 📈 Dashboard: Phân tích điện áp theo Zone_Bx, có biểu đồ và xuất báo cáo Excel/Word\n"
            "   • 📤 Xuất TBA lỗi: Xuất danh sách trạm chưa ánh xạ Zone_Bx ra file Excel\n\n"
            "2. Bộ lọc dữ liệu:\n"
            "   • Lọc theo Trạm biến áp (gõ tên trạm)\n"
            "   • Lọc theo U danh định (Uđd)\n"
            "   • Lọc theo Zone_Bx\n"
            "   • Lọc theo Thời gian (từ ngày – đến ngày)\n"
            "   • Lọc theo ngưỡng: U THẤP (≤ %Uđd), U CAO (≥ %Uđd)\n\n"
            "3. Biểu đồ phân tích:\n"
            "   • Line / Scatter U thực tế\n"
            "   • 🌡 Heatmap điện áp theo giờ/ngày\n"
            "   • 📊 Histogram phân phối U\n"
            "   • 📦 Boxplot U\n"
            "   • 💾 Lưu hình: Xuất biểu đồ ra PNG\n\n"
            "4. Khác:\n"
            "   • Bộ nhớ cache: tự động lưu dữ liệu, có thể xoá toàn bộ khi cần\n"
            "   • Dashboard có thể xuất báo cáo Excel / Word để in ấn & chia sẻ\n\n"
            "5. LƯU Ý:\n"
            "   • Phải có file Excel DB_VietSub.xlsx để tham chiếu, tìm Zone_Bx\n\n"
            "——————————————\n"
            "Bản quyền phần mềm © 2025 NSO / SuNV\n"
            "Liên hệ hỗ trợ: 0966 736 889"
        )

        textbox = ctk.CTkTextbox(
            win, width=560, height=370,
            font=("Segoe UI", 13), wrap="word"
        )
        textbox.insert("1.0", help_text)
        textbox.configure(state="disabled")  # chỉ đọc
        textbox.pack(padx=20, pady=5, fill="both", expand=True)

        ctk.CTkButton(win, text="Đóng", command=win.destroy).pack(pady=12)
    def _kpi_card(self, parent, icon, label, value, color, col):
        card = ctk.CTkFrame(parent, fg_color=color, corner_radius=12, width=104, height=64)
        card.grid(row=0, column=col, padx=12, pady=0, sticky="nsew")
        ctk.CTkLabel(card, text=icon, font=("Segoe UI", 23)).pack(side="top", pady=(6, 0))

        # value có thể là StringVar để cập nhật live
        if hasattr(value, "get") and hasattr(value, "set"):
            ctk.CTkLabel(card, textvariable=value, font=("Segoe UI", 18, "bold"),
                         text_color="#fff").pack(side="top", pady=(0, 1))
        else:
            ctk.CTkLabel(card, text=str(value), font=("Segoe UI", 18, "bold"),
                         text_color="#fff").pack(side="top", pady=(0, 1))

        ctk.CTkLabel(card, text=label, font=("Segoe UI", 11, "bold"),
                     text_color="#fff").pack(side="top", pady=(0, 4))


# ==================== Entrypoint ====================
def main():
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()
